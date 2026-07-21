"""Tests for the FS Review deterministic engine.

Run locally:  pip install -r requirements.txt pytest && pytest -q
Runs automatically on every GitHub push via .github/workflows/ci.yml.
"""
import io
import os
import importlib.util

import pytest
from docx import Document

os.environ.pop("DATABASE_URL", None)          # file-fallback mode for tests

HERE = os.path.dirname(os.path.abspath(__file__))
spec = importlib.util.spec_from_file_location("fsapp", os.path.join(HERE, "app.py"))
m = importlib.util.module_from_spec(spec)
spec.loader.exec_module(m)


# ---------------------------------------------------------------- helpers
def make_doc(tables):
    """Build an in-memory .docx with the given tables (list of list-of-rows)."""
    doc = Document()
    for rows in tables:
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                t.rows[r].cells[c].text = str(val)
    return doc


BS = [  # simple balance sheet, deliberately out by 1
    ["", "2025", "2024"],
    ["Cash and cash equivalents", "1,000", "800"],
    ["Total Current Assets", "1,000", "800"],
    ["Total Assets", "1,000", "800"],
    ["Amount due to director", "400", "300"],
    ["Total Current Liabilities", "400", "300"],
    ["Share capital", "1", "1"],
    ["Retained earnings", "600", "499"],
    ["Total Equity", "601", "500"],
    ["Total Equity and Liabilities", "1,001", "800"],
]

PL = [
    ["", "2025", "2024"],
    ["Revenue", "5,000", "4,000"],
    ["Cost of sales", "(2,000)", "(1,500)"],
    ["Gross profit", "3,000", "2,500"],
    ["Administrative expenses", "(1,000)", "(900)"],
    ["Profit before tax", "2,000", "1,600"],
    ["Income tax", "(340)", "(272)"],
    ["Profit for the year", "1,660", "1,328"],
]


# ---------------------------------------------------------------- tests
def test_balance_equation_flags_mismatch():
    doc = make_doc([BS])
    res = m.check_balance_equation(doc)
    assert any(not r["balanced"] and abs(r["difference"]) == 1.0 for r in res)


def test_pl_casts_clean():
    doc = make_doc([PL])
    issues = [c for i, t in enumerate(doc.tables) for c in m.check_pl(i, t)]
    assert issues == []


def test_pl_flags_wrong_profit():
    bad = [row[:] for row in PL]
    bad[7] = ["Profit for the year", "1,700", "1,328"]   # should be 1,660
    doc = make_doc([bad])
    issues = [c for i, t in enumerate(doc.tables) for c in m.check_pl(i, t)]
    assert any(abs(c["difference"]) == 40.0 for c in issues)


def test_related_party_flags_director_loan():
    doc = make_doc([BS])
    flags = m.check_related_party_loans(doc)
    assert any("due to director" in f["error"].lower() for f in flags)


def test_note_face_tie_flags_mismatch():
    note = [["", "2025", "2024"], ["Cash and cash equivalents", "999", "800"]]
    doc = make_doc([BS, note])
    ties = m.check_note_face_ties(doc)
    assert any("Cash" in t["check"] for t in ties)


def test_note_face_tie_clean_when_totals_match():
    note = [["", "2025", "2024"], ["Cash and cash equivalents", "1,000", "800"]]
    doc = make_doc([BS, note])
    assert m.check_note_face_ties(doc) == []


def test_language_flags_phrasing_and_typos():
    doc = Document()
    doc.add_paragraph("The Company are engaged in trading. Management yearically evaluates.")
    issues = m.check_language(doc)
    found = {i["found"] for i in issues}
    assert "The Company are" in found and "yearically" in found


def test_corrections_are_labelled_and_deduped():
    findings = {
        "balance_checks": [{"balanced": False, "total_assets": 1000.0,
                            "equity_plus_liabilities": 1001.0, "difference": -1.0}],
        "pl_checks": [], "tally_checks": [], "row_checks": [], "cross_checks": [],
        "going_concern": {}, "acra": {}, "language_issues": [],
        "related_party": [], "ai": {},
    }
    C = m.build_corrections(findings)
    assert C and C[0]["error"].startswith("ARITHMETIC")
    assert "Difference to correct" in C[0]["error"]


def test_disclosure_templates_trigger_on_director_loan():
    findings = {"related_party": [{"severity": "high", "kind": "loan to director",
                                   "error": "x", "recommendation": "y"}],
                "pl_checks": [], "frs_checks": [], "going_concern": {}}
    T = m.build_disclosure_templates(findings, "dividend declared")
    titles = " ".join(t["title"] for t in T)
    assert "s.162" in titles and "Related party" in titles and "Dividends" in titles


def test_marked_fs_has_comments_and_valid_docx():
    doc = make_doc([BS])
    buf = io.BytesIO()
    doc.save(buf)
    findings = {
        "corrections": [{"severity": "high",
                         "error": "ARITHMETIC — statement of financial position does not "
                                  "balance. · Total assets stated: 1,000.00 · Total equity "
                                  "+ liabilities: 1,001.00 · Difference to correct: 1.00",
                         "recommendation": "Recheck the subtotals."}],
        "language_issues": [], "ai": {}, "pl_checks": [], "tally_checks": [],
        "row_checks": [], "disclosure_templates": [],
    }
    out = m.build_marked_fs({"findings": findings}, buf.getvalue())
    import zipfile
    z = zipfile.ZipFile(out)
    assert "word/comments.xml" in z.namelist()
    cx = z.read("word/comments.xml").decode()
    assert "Review point 1" in cx and "Difference to correct" in cx
    out.seek(0)
    Document(out)          # python-docx must be able to reopen the package
