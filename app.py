"""
FS Review portal
================
A small, secure Flask web app for uploading financial-statement files
(.docx / .pdf / .xlsx) and viewing an automated review report.

- Login required (passwords are hashed, never stored in plain text).
- Sessions are signed with a secret key.
- Uploaded files are stored on disk; a lightweight automated review runs
  on .docx files and the findings are shown as a report.

Run locally:    python app.py
Production:     gunicorn app:app   (see README.md)
"""

import os
import io
import re
import json
import uuid
import secrets
import datetime as dt
from functools import wraps

from flask import (
    Flask, render_template, request, redirect, url_for,
    session, flash, send_from_directory, abort, Response,
)
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename

# --------------------------------------------------------------------------
# Configuration
# --------------------------------------------------------------------------
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
DATA_DIR = os.path.join(BASE_DIR, "data")
USERS_FILE = os.path.join(DATA_DIR, "users.json")
RECORDS_FILE = os.path.join(DATA_DIR, "records.json")
AUDIT_FILE = os.path.join(DATA_DIR, "audit_log.json")

ALLOWED_EXTENSIONS = {".docx", ".pdf", ".xlsx", ".xls"}
MAX_CONTENT_LENGTH = 25 * 1024 * 1024  # 25 MB

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(DATA_DIR, exist_ok=True)

app = Flask(__name__)
# Secret key: set FLASK_SECRET_KEY in production. Falls back to a generated
# value for local use (sessions reset on restart if not set).
app.config["SECRET_KEY"] = os.environ.get("FLASK_SECRET_KEY", os.urandom(32).hex())
app.config["MAX_CONTENT_LENGTH"] = MAX_CONTENT_LENGTH
# Session cookie hardening (HTTPS-only, no JS access, CSRF-resistant).
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
app.config["SESSION_COOKIE_SECURE"] = os.environ.get("COOKIE_SECURE", "1") == "1"
app.config["PERMANENT_SESSION_LIFETIME"] = 60 * 60 * 8  # 8 hours

# Optional persistent database (Neon/Supabase Postgres). When DATABASE_URL is
# set, users + reviewed records + uploaded files are stored there so they
# survive restarts/redeploys. Without it, the app falls back to JSON files.
DATABASE_URL = os.environ.get("DATABASE_URL")
USE_DB = bool(DATABASE_URL)

# Passwordless magic-link login (email a one-time link). Active only when a
# database is connected (to store the one-time tokens) AND a Resend API key is
# set (to send the email). Otherwise the app stays on username/password login.
RESEND_API_KEY = os.environ.get("RESEND_API_KEY")
MAIL_FROM = os.environ.get("MAIL_FROM", "FS Review <onboarding@resend.dev>")
ADMIN_EMAIL = os.environ.get("ADMIN_EMAIL", "").strip().lower()
MAGIC_LOGIN = USE_DB and bool(RESEND_API_KEY)
TOKEN_TTL_MIN = 20  # magic-link validity, minutes


# --------------------------------------------------------------------------
# Templates are embedded here (no separate templates/ folder needed)
# --------------------------------------------------------------------------
from jinja2 import DictLoader

BASE_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{% block title %}FS Review{% endblock %}</title>
  <style>
    :root{
      --bg:#f4f6f9; --card:#ffffff; --ink:#1f2937; --muted:#6b7280;
      --line:#e5e7eb; --brand:#1d4ed8; --brand-dark:#1e40af;
      --good:#047857; --good-bg:#ecfdf5; --bad:#b91c1c; --bad-bg:#fef2f2;
      --warn:#92400e; --warn-bg:#fffbeb;
    }
    *{box-sizing:border-box}
    body{margin:0;font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,Helvetica,Arial,sans-serif;
      background:var(--bg);color:var(--ink);line-height:1.5}
    a{color:var(--brand);text-decoration:none}
    a:hover{text-decoration:underline}
    .topbar{background:var(--brand);color:#fff;padding:14px 24px;display:flex;
      align-items:center;justify-content:space-between}
    .topbar .brand{font-weight:700;font-size:18px;letter-spacing:.2px}
    .topbar a{color:#dbeafe}
    .wrap{max-width:920px;margin:32px auto;padding:0 20px}
    .card{background:var(--card);border:1px solid var(--line);border-radius:12px;
      padding:24px;box-shadow:0 1px 3px rgba(0,0,0,.04);margin-bottom:20px}
    h1{font-size:22px;margin:0 0 4px}
    h2{font-size:16px;margin:0 0 12px;color:var(--ink)}
    .muted{color:var(--muted);font-size:14px}
    .btn{display:inline-block;background:var(--brand);color:#fff;border:none;
      padding:10px 18px;border-radius:8px;font-size:15px;cursor:pointer;font-weight:600}
    .btn:hover{background:var(--brand-dark);text-decoration:none}
    .btn.secondary{background:#fff;color:var(--brand);border:1px solid var(--brand)}
    .btn.danger{background:#fff;color:var(--bad);border:1px solid var(--bad);padding:6px 12px;font-size:13px}
    input[type=text],input[type=password]{width:100%;padding:11px 12px;border:1px solid var(--line);
      border-radius:8px;font-size:15px;margin-top:6px}
    label{font-size:14px;font-weight:600}
    table{width:100%;border-collapse:collapse;font-size:14px}
    th,td{text-align:left;padding:10px 8px;border-bottom:1px solid var(--line)}
    th{color:var(--muted);font-weight:600;font-size:12px;text-transform:uppercase;letter-spacing:.4px}
    .flash{padding:12px 16px;border-radius:8px;margin-bottom:16px;font-size:14px}
    .flash.error{background:var(--bad-bg);color:var(--bad);border:1px solid #fecaca}
    .flash.success{background:var(--good-bg);color:var(--good);border:1px solid #a7f3d0}
    .pill{display:inline-block;padding:3px 10px;border-radius:999px;font-size:12px;font-weight:600}
    .pill.good{background:var(--good-bg);color:var(--good)}
    .pill.bad{background:var(--bad-bg);color:var(--bad)}
    .pill.warn{background:var(--warn-bg);color:var(--warn)}
    .dropzone{border:2px dashed #c7d2fe;border-radius:12px;padding:28px;text-align:center;
      background:#f8faff;margin:8px 0 16px}
    ul{margin:8px 0;padding-left:20px}
    li{margin:3px 0}
  </style>
</head>
<body>
  <div class="topbar">
    <div class="brand">{% if has_logo %}<img src="{{ url_for('logo') }}" alt="Assembly Works" style="height:26px;vertical-align:middle;margin-right:10px;background:#fff;border-radius:4px;padding:2px 6px">{% endif %}ASSEMBLY WORKS <span style="font-weight:400;opacity:.85">· FS Review Portal</span></div>
    <div>
      {% if session.get('user') %}
        {% if current_is_admin %}<a href="{{ url_for('admin_audit') }}" style="margin-right:14px">Audit log</a>{% endif %}
        {% if current_user_can_invite %}<a href="{{ url_for('admin_users') }}" style="margin-right:14px">People</a>{% endif %}
        <span style="color:#dbeafe;font-size:14px">{{ session.get('name') }}</span>
        &nbsp;·&nbsp; <a href="{{ url_for('logout') }}">Log out</a>
      {% endif %}
    </div>
  </div>
  <div class="wrap">
    {% with messages = get_flashed_messages(with_categories=true) %}
      {% for category, message in messages %}
        <div class="flash {{ category }}">{{ message }}</div>
      {% endfor %}
    {% endwith %}
    {% block content %}{% endblock %}
  </div>
</body>
</html>"""

LOGIN_HTML = """{% extends "base.html" %}
{% block title %}Sign in · FS Review{% endblock %}
{% block content %}
<div class="card" style="max-width:420px;margin:40px auto">
  <h1>Sign in</h1>
  <p class="muted">Sign in with your email and password. Only invited people can sign in.</p>
  <form method="post" action="{{ url_for('login') }}" style="margin-top:16px">
    <div style="margin-bottom:14px">
      <label for="email">Email</label>
      <input type="text" id="email" name="email" autocomplete="username" required autofocus>
    </div>
    <div style="margin-bottom:20px">
      <label for="password">Password</label>
      <input type="password" id="password" name="password" autocomplete="current-password">
    </div>
    <button class="btn" type="submit" style="width:100%">Sign in</button>
    {% if magic %}
    <button class="btn secondary" type="submit" name="sendotp" value="1"
            style="width:100%;margin-top:10px">Email me a one-time code (OTP)</button>
    <button class="btn secondary" type="submit" name="sendlink" value="1"
            style="width:100%;margin-top:10px">Email me a sign-in link instead</button>
    {% endif %}
  </form>
</div>
{% endblock %}"""

OTP_HTML = """{% extends "base.html" %}
{% block title %}Enter code · FS Review{% endblock %}
{% block content %}
<div class="card" style="max-width:420px;margin:40px auto">
  <h1>Enter your sign-in code</h1>
  <p class="muted">We emailed a 6-digit code to <strong>{{ email }}</strong>. It expires in 10 minutes.</p>
  <form method="post" action="{{ url_for('otp_verify') }}" style="margin-top:16px">
    <input type="hidden" name="email" value="{{ email }}">
    <div style="margin-bottom:20px">
      <label for="code">6-digit code</label>
      <input type="text" id="code" name="code" inputmode="numeric" pattern="[0-9 ]*"
             maxlength="7" autocomplete="one-time-code" required autofocus
             style="font-size:24px;letter-spacing:8px;text-align:center">
    </div>
    <button class="btn" type="submit" style="width:100%">Sign in</button>
  </form>
  <form method="post" action="{{ url_for('login') }}" style="margin-top:10px">
    <input type="hidden" name="email" value="{{ email }}">
    <button class="btn secondary" type="submit" name="sendotp" value="1"
            style="width:100%">Send a new code</button>
  </form>
  <p class="muted" style="margin-top:12px"><a href="{{ url_for('login') }}">← Back to sign in</a></p>
</div>
{% endblock %}"""

DASHBOARD_HTML = """{% extends "base.html" %}
{% block title %}Dashboard · FS Review{% endblock %}
{% block content %}
<div class="card">
  <h1>Upload financial statements</h1>
  <p class="muted">Allowed file types: .docx, .pdf, .xlsx, .xls — max 25 MB. A review report is generated on upload.</p>
  <form method="post" action="{{ url_for('upload') }}" enctype="multipart/form-data" id="uploadForm">
    <div class="dropzone" id="dz">
      <p class="muted" style="margin:0 0 10px">Drag &amp; drop your financial statements here, or choose a file:</p>
      <input type="file" id="fileInput" name="file" accept=".docx,.pdf,.xlsx,.xls" required
             style="font-size:15px">
    </div>
    <script>
    document.addEventListener('DOMContentLoaded', function(){
      function wire(zoneId, inputId){
        var dz = document.getElementById(zoneId), fi = document.getElementById(inputId);
        if(!dz || !fi) return;
        ['dragover','dragenter'].forEach(function(ev){
          dz.addEventListener(ev, function(e){ e.preventDefault();
            dz.style.background = '#eef2ff'; dz.style.borderColor = '#1d4ed8'; });
        });
        ['dragleave','drop'].forEach(function(ev){
          dz.addEventListener(ev, function(e){ e.preventDefault();
            dz.style.background = ''; dz.style.borderColor = ''; });
        });
        dz.addEventListener('drop', function(e){
          e.preventDefault();
          if(e.dataTransfer && e.dataTransfer.files.length){ fi.files = e.dataTransfer.files; }
        });
      }
      wire('dz', 'fileInput');
      wire('dzb', 'bizInput');
      // Prevent double-submit (which created duplicate reviews) and show progress.
      var uf = document.getElementById('uploadForm'),
          ub = document.getElementById('uploadBtn'),
          fin = document.getElementById('fileInput');
      if(uf && ub){
        uf.addEventListener('submit', function(){
          if(fin && (!fin.files || !fin.files.length)) return;   // let 'required' handle it
          ub.disabled = true;
          ub.textContent = 'Reviewing… this can take up to a minute — please wait';
        });
      }
    });
    </script>
    <div style="margin:0 0 16px">
      <label style="font-size:13px">Latest ACRA BizFile (Business Profile PDF) <span class="muted">(optional — crawled to cross-check UEN, share capital &amp; directors)</span></label>
      <div class="dropzone" id="dzb" style="margin-top:6px">
        <p class="muted" style="margin:0 0 10px">Drag &amp; drop the ACRA BizFile PDF here, or choose a file:</p>
        <input type="file" id="bizInput" name="acra_bizfile" accept=".pdf" style="font-size:15px">
      </div>
    </div>
    <button class="btn" type="submit" id="uploadBtn">Upload &amp; review</button>
  </form>
</div>

<div class="card">
  <h2>Reviewed files</h2>
  {% if records %}
  <table>
    <thead>
      <tr><th>File</th><th>Type</th><th>Size</th><th>Uploaded</th><th></th></tr>
    </thead>
    <tbody>
      {% for r in records %}
      <tr>
        <td><a href="{{ url_for('report', rec_id=r.id) }}">{{ r.original_name }}</a><br>
            <span class="muted" style="font-size:12px">by {{ r.uploaded_by }}</span></td>
        <td>{{ r.ext.lstrip('.')|upper }}</td>
        <td>{{ r.size_bytes|filesize }}</td>
        <td class="muted">{{ r.uploaded_at.replace('T',' ') }}</td>
        <td><a href="{{ url_for('report', rec_id=r.id) }}">View report →</a></td>
      </tr>
      {% endfor %}
    </tbody>
  </table>
  {% else %}
  <p class="muted">No files reviewed yet. Upload one above to get started.</p>
  {% endif %}
</div>
{% endblock %}"""

REPORT_HTML = """{% extends "base.html" %}
{% block title %}Review report · {{ r.original_name }}{% endblock %}
{% block content %}
<div class="card">
  <p class="muted"><a href="{{ url_for('dashboard') }}">← Back to dashboard</a></p>
  <h1>Review report</h1>
  <p class="muted">{{ r.original_name }} · {{ r.ext.lstrip('.')|upper }} · {{ r.size_bytes|filesize }}
     · uploaded {{ r.uploaded_at.replace('T',' ') }} by {{ r.uploaded_by }}</p>
  <p style="margin-top:14px">
    <a class="btn secondary" href="{{ url_for('download', rec_id=r.id) }}">Download original</a>
    <a class="btn" href="{{ url_for('download_report', rec_id=r.id) }}">Download review report (Word)</a>
    {% if r.ext == '.docx' %}<a class="btn" href="{{ url_for('download_revised', rec_id=r.id) }}">Download revised FS (marked up)</a>{% endif %}
  </p>
</div>

{% set f = r.findings %}

{% if f.error %}
<div class="card">
  <span class="pill bad">Could not parse</span>
  <p style="margin-top:12px">{{ f.error }}</p>
</div>
{% else %}

{% set hi = f.corrections|selectattr('severity','equalto','high')|list|length %}
{% set md = f.corrections|selectattr('severity','equalto','medium')|list|length %}
{% set lo = f.corrections|length - hi - md %}
<div class="card" style="border-left:6px solid {% if hi %}#dc2626{% elif md %}#d97706{% else %}#16a34a{% endif %}">
  <h2 style="margin-top:0">At a glance</h2>
  <p style="font-size:15px;margin:6px 0">
    <span class="pill bad">{{ hi }} High</span>
    <span class="pill warn">{{ md }} Medium</span>
    <span class="pill">{{ lo }} Minor</span>
    &nbsp; {% if f.balance_checks and f.balance_checks[0].balanced %}<span class="pill good">Balance sheet balances</span>{% endif %}
    {% if f.going_concern and f.going_concern.at_risk %}<span class="pill bad">Going-concern risk</span>{% else %}<span class="pill good">No going-concern flag</span>{% endif %}
    {% if f.ai and f.ai.enabled %}<span class="pill good">AI reviewed</span>{% else %}<span class="pill warn">AI off</span>{% endif %}
  </p>
  {% if f.ai and f.ai.narrative %}<p class="muted" style="margin:8px 0 0">{{ f.ai.narrative }}</p>{% endif %}
  <p class="muted" style="margin:10px 0 0;font-size:12px">Work top-down: fix the High items first. Full detail is in the sections below.</p>
</div>

{% if f.corrections %}
<div class="card" style="border-color:#fca5a5;background:#fff5f5">
  <h2>Corrections to make <span class="pill bad">{{ f.corrections|length }}</span></h2>
  <p class="muted">A punch-list to work through before finalising. Each row shows the error found and the recommended correction. Arithmetic and disclosure points come from the automated checks; judgement points come from the AI reviewer.</p>
  <table style="margin-top:10px">
    <thead><tr><th style="width:70px">Severity</th><th style="width:44px">#</th><th>Error found</th><th>Recommended correction</th></tr></thead>
    <tbody>
    {% for c in f.corrections %}
    <tr>
      <td>{% if c.severity == 'high' %}<span class="pill bad">High</span>{% elif c.severity == 'medium' %}<span class="pill warn">Medium</span>{% else %}<span class="pill">Minor</span>{% endif %}</td>
      <td>{{ loop.index }}</td>
      <td>{{ c.error }}</td>
      <td>{{ c.recommendation }}</td>
    </tr>
    {% endfor %}
    </tbody>
  </table>
</div>
{% endif %}

<div class="card">
  <h2>Document structure</h2>
  <p><strong>{{ f.tables }}</strong> table(s), <strong>{{ f.paragraph_count }}</strong> non-empty paragraph(s).</p>
  {% if f.sections_found %}
    <p style="margin-top:12px"><strong>Sections detected:</strong></p>
    <ul>{% for s in f.sections_found %}<li>{{ s }}</li>{% endfor %}</ul>
  {% endif %}
  {% if f.sections_missing %}
    <p style="margin-top:12px"><strong>Common sections not detected</strong>
       <span class="muted">(may be absent, named differently, or in a separate file):</span></p>
    <ul>{% for s in f.sections_missing %}<li class="muted">{{ s }}</li>{% endfor %}</ul>
  {% endif %}
</div>

<div class="card">
  <h2>ACRA verification
    {% if not f.acra.enabled %}<span class="pill warn">no UEN found</span>
    {% elif f.acra.found and f.acra.name_matches %}<span class="pill good">UEN verified</span>
    {% elif f.acra.found %}<span class="pill warn">check name</span>
    {% else %}<span class="pill bad">UEN not in register</span>{% endif %}</h2>
  {% if f.acra.error and not f.acra.found %}<p class="muted">{{ f.acra.error }}</p>{% endif %}
  {% if f.acra.uen %}
  <table>
    <tbody>
      <tr><td style="width:38%"><strong>UEN in document</strong></td><td>{{ f.acra.uen }}</td></tr>
      {% if f.acra.found %}
      <tr><td><strong>ACRA registered name</strong></td><td>{{ f.acra.official_name }}
        {% if f.acra.name_matches %}<span class="pill good">matches</span>{% else %}<span class="pill bad">differs from document</span>{% endif %}</td></tr>
      <tr><td><strong>Registration status</strong></td><td>{{ f.acra.status }}</td></tr>
      <tr><td><strong>Registered address</strong></td><td>{{ f.acra.address }}</td></tr>
      {% endif %}
      <tr><td><strong>Issued &amp; paid-up share capital (per FS)</strong></td><td>{% if f.acra.fs_share_capital is not none %}${{ "{:,.2f}".format(f.acra.fs_share_capital) }}{% else %}not detected in the accounts{% endif %}</td></tr>
      {% if f.acra.registered_share_capital is not none %}
      <tr><td><strong>Issued &amp; paid-up share capital (per ACRA)</strong></td><td>${{ "{:,.2f}".format(f.acra.registered_share_capital) }}
        {% if f.acra.fs_share_capital is none %}<span class="pill warn">FS figure not detected — compare manually</span>
        {% elif f.acra.share_capital_matches %}<span class="pill good">agrees with FS</span>
        {% else %}<span class="pill bad">differs from FS by ${{ "{:,.2f}".format((f.acra.registered_share_capital - f.acra.fs_share_capital)|abs) }}</span>{% endif %}</td></tr>
      {% if f.acra.fs_share_capital is not none and not f.acra.share_capital_matches %}
      <tr><td></td><td class="muted">Share capital in the accounts should equal the issued &amp; paid-up capital registered at ACRA. Reconcile before finalising — common causes are the FS showing the <em>number of shares</em> rather than the dollar amount, a share allotment made after the financial year-end, capital shown net of issue costs, or a transposition error.</td></tr>
      {% endif %}
      {% endif %}
    </tbody>
  </table>
  {% endif %}
  {% if f.acra.bizfile %}
    {% set bz = f.acra.bizfile %}
    {% if bz.error %}<p class="muted" style="margin-top:10px">BizFile: {{ bz.error }}</p>{% endif %}
    {% if bz.entity_name or bz.status or bz.shareholders or bz.directors or bz.issued_share_capital is not none %}
    <p style="margin-top:12px"><strong>Crawled from the uploaded ACRA BizFile</strong></p>
    <table><tbody>
      {% if bz.entity_name %}<tr><td style="width:38%"><strong>Name</strong></td><td>{{ bz.entity_name }}</td></tr>{% endif %}
      {% if bz.status %}<tr><td><strong>Status</strong></td><td>{{ bz.status }}</td></tr>{% endif %}
      {% if bz.issued_share_capital is not none %}<tr><td><strong>Issued share capital</strong></td><td>{{ "{:,.2f}".format(bz.issued_share_capital) }}</td></tr>{% endif %}
      {% if bz.paid_up_capital is not none %}<tr><td><strong>Paid-up capital</strong></td><td>{{ "{:,.2f}".format(bz.paid_up_capital) }}</td></tr>{% endif %}
      {% if bz.shareholders %}<tr><td><strong>Shareholders</strong></td><td>{% for s in bz.shareholders %}{{ s.name }}{% if s.shares %} ({{ s.shares }}){% endif %}{% if not loop.last %}; {% endif %}{% endfor %}</td></tr>{% endif %}
      {% if bz.directors %}<tr><td><strong>Directors</strong></td><td>{% for d in bz.directors %}{{ d.name }}{% if not loop.last %}; {% endif %}{% endfor %}</td></tr>{% endif %}
    </tbody></table>
    {% endif %}
  {% else %}
  <p class="muted" style="margin-top:10px">Share capital is not in ACRA's free data. Upload the latest ACRA BizFile (Business Profile PDF) on the dashboard to cross-check share capital and shareholders.</p>
  {% endif %}
</div>

<div class="card">
  <h2>Arithmetic / tally checks {% if f.tally_checks %}<span class="pill bad">{{ f.tally_checks|length }} flagged</span>{% else %}<span class="pill good">No mismatches found</span>{% endif %}</h2>
  {% if f.tally_checks %}
  <p class="muted">Subtotals/totals where the lines above do not add up to the stated figure — verify each:</p>
  <table>
    <thead><tr><th>Table</th><th>Total line</th><th>Issue</th><th>Sum of lines</th><th>Stated</th><th>Difference</th></tr></thead>
    <tbody>
      {% for c in f.tally_checks %}
      <tr>
        <td>{{ c.table }}</td><td>{{ c.label }}</td>
        <td>{% if c.kind == 'sign / brackets' %}<span class="pill warn">sign / brackets</span>{% else %}<span class="pill bad">does not add up</span>{% endif %}</td>
        <td>{{ "{:,.2f}".format(c.sum_of_parts) }}</td>
        <td>{{ "{:,.2f}".format(c.stated_total) }}</td>
        <td style="color:#b91c1c">{{ "{:,.2f}".format(c.difference) }}</td>
      </tr>
      {% endfor %}
    </tbody>
  </table>
  {% else %}
  <p class="muted">No column-total mismatches were detected.</p>
  {% endif %}
</div>

<div class="card">
  <h2>Balance sheet equation
    {% if f.balance_checks %}
      {% set unbalanced = f.balance_checks | selectattr('balanced', 'equalto', false) | list %}
      {% if unbalanced %}<span class="pill bad">Does not balance</span>{% else %}<span class="pill good">Balances</span>{% endif %}
    {% else %}<span class="pill warn">Not found</span>{% endif %}</h2>
  {% if f.balance_checks %}
  <p class="muted">Total assets vs. total equity + total liabilities:</p>
  <table>
    <thead><tr><th>Total assets</th><th>Equity + liabilities</th><th>Difference</th><th>Status</th></tr></thead>
    <tbody>
      {% for b in f.balance_checks %}
      <tr>
        <td>{{ "{:,.2f}".format(b.total_assets) }}</td>
        <td>{{ "{:,.2f}".format(b.equity_plus_liabilities) }}</td>
        <td {% if not b.balanced %}style="color:#b91c1c"{% endif %}>{{ "{:,.2f}".format(b.difference) }}</td>
        <td>{% if b.balanced %}<span class="pill good">OK</span>{% else %}<span class="pill bad">Off</span>{% endif %}</td>
      </tr>
      {% endfor %}
    </tbody>
  </table>
  {% else %}
  <p class="muted">Could not locate "Total assets" / "Total equity" / "Total liabilities" lines to test the balance equation. Check the labels in the statement of financial position.</p>
  {% endif %}
</div>

<div class="card">
  <h2>Profit &amp; loss flow {% if f.pl_checks %}<span class="pill bad">{{ f.pl_checks|length }} flagged</span>{% else %}<span class="pill good">Consistent</span>{% endif %}</h2>
  {% if f.pl_checks %}
  <table>
    <thead><tr><th>Table</th><th>Check</th><th>Expected</th><th>Stated</th><th>Difference</th></tr></thead>
    <tbody>
      {% for c in f.pl_checks %}
      <tr><td>{{ c.table }}</td><td>{{ c.check }}</td>
        <td>{{ "{:,.2f}".format(c.expected) }}</td>
        <td>{{ "{:,.2f}".format(c.stated) }}</td>
        <td style="color:#b91c1c">{{ "{:,.2f}".format(c.difference) }}</td></tr>
      {% endfor %}
    </tbody>
  </table>
  {% else %}
  <p class="muted">Gross profit and loss/profit for the year tie to their components (or no P&amp;L detected).</p>
  {% endif %}
</div>

<div class="card">
  <h2>Cross-add checks (changes in equity, etc.) {% if f.row_checks %}<span class="pill bad">{{ f.row_checks|length }} flagged</span>{% else %}<span class="pill good">OK</span>{% endif %}</h2>
  {% if f.row_checks %}
  <table>
    <thead><tr><th>Table</th><th>Row</th><th>Sum across</th><th>Stated total</th><th>Difference</th></tr></thead>
    <tbody>
      {% for c in f.row_checks %}
      <tr><td>{{ c.table }}</td><td>{{ c.row }}</td>
        <td>{{ "{:,.2f}".format(c.sum_across) }}</td>
        <td>{{ "{:,.2f}".format(c.stated_total) }}</td>
        <td style="color:#b91c1c">{{ "{:,.2f}".format(c.difference) }}</td></tr>
      {% endfor %}
    </tbody>
  </table>
  {% else %}
  <p class="muted">Rows with a "Total" column add across correctly (or none found).</p>
  {% endif %}
</div>

{% if f.cross_checks is defined %}
<div class="card">
  <h2>Cross-statement ties {% if f.cross_checks %}<span class="pill bad">{{ f.cross_checks|length }} flagged</span>{% else %}<span class="pill good">Tie up</span>{% endif %}</h2>
  {% if f.cross_checks %}
  <table>
    <thead><tr><th>Check</th><th>Computed</th><th>Stated</th><th>Difference</th></tr></thead>
    <tbody>
      {% for c in f.cross_checks %}
      <tr><td>{{ c.check }}</td>
        <td>{{ "{:,.2f}".format(c.left) }}</td>
        <td>{{ "{:,.2f}".format(c.right) }}</td>
        <td style="color:#b91c1c">{{ "{:,.2f}".format(c.difference) }}</td></tr>
      {% endfor %}
    </tbody>
  </table>
  {% else %}
  <p class="muted">Cash-flow closing cash ties to opening plus net movement (or the lines weren't located).</p>
  {% endif %}
</div>
{% endif %}

<div class="card">
  <h2>FRS disclosure indicators</h2>
  <p class="muted">A keyword scan for common Singapore FRS disclosures — an <em>indicator</em>, not a compliance conclusion. "Not detected" means the wording wasn't found and should be confirmed manually; the AI reviewer below assesses adequacy in context.</p>
  <table>
    <thead><tr><th>FRS</th><th>Disclosure</th><th>Detected?</th></tr></thead>
    <tbody>
      {% for k in f.frs_checks %}
      <tr>
        <td>{{ k.frs }}</td><td>{{ k.item }}</td>
        <td>{% if k.present %}<span class="pill good">Detected</span>{% else %}<span class="pill warn">Not detected — confirm</span>{% endif %}</td>
      </tr>
      {% endfor %}
    </tbody>
  </table>
</div>

<div class="card">
  {% set gc = f.going_concern %}
  <h2>Going concern (FRS 1)
    {% if gc.verdict_level == 'bad' %}<span class="pill bad">risk — disclosure required</span>
    {% elif gc.verdict_level == 'warn' %}<span class="pill warn">review</span>
    {% elif gc.verdict_level == 'good' %}<span class="pill good">solvent</span>
    {% endif %}</h2>

  {% if gc.verdict and (gc.equity is not none or gc.net_current is not none) %}
  <p class="muted">Solvency read from the statement of financial position:</p>
  <table>
    <thead><tr><th>Indicator</th><th>Amount</th><th>Read</th></tr></thead>
    <tbody>
      {% if gc.equity is not none %}
      <tr><td>Total equity (net assets)</td><td>{{ "{:,.2f}".format(gc.equity) }}</td>
          <td>{% if gc.equity < 0 %}<span class="pill bad">Net liabilities</span>{% else %}<span class="pill good">Positive</span>{% endif %}</td></tr>
      {% endif %}
      {% if gc.net_current is not none %}
      <tr><td>Net current position</td><td>{{ "{:,.2f}".format(gc.net_current) }}</td>
          <td>{% if gc.net_current < 0 %}<span class="pill bad">Net current liabilities</span>{% else %}<span class="pill good">Net current assets</span>{% endif %}</td></tr>
      {% endif %}
    </tbody>
  </table>
  {% endif %}

  {% if gc.verdict %}<p style="margin-top:12px">{{ gc.verdict }}</p>{% endif %}

  {% if gc.contradictions %}
  <ul>{% for c in gc.contradictions %}<li style="color:#b91c1c">{{ c }}</li>{% endfor %}</ul>
  {% endif %}

  {% if gc.elements %}
  <p style="margin-top:12px" class="muted">Disclosure checklist{% if gc.at_risk %} — enhanced going-concern note expected{% endif %}:</p>
  <table>
    <thead><tr><th>Element</th><th>In the note?</th></tr></thead>
    <tbody>
      {% for e in gc.elements %}
      <tr><td>{{ e.element }}</td><td>{% if e.present %}<span class="pill good">Yes</span>{% else %}<span class="pill warn">Check</span>{% endif %}</td></tr>
      {% endfor %}
    </tbody>
  </table>
  {% endif %}
</div>

<div class="card">
  <h2>AI reviewer — FRS judgement &amp; drafting
    {% if f.ai.enabled %}{% if f.ai.frs_observations %}<span class="pill warn">{{ f.ai.frs_observations|length }} observation(s)</span>{% else %}<span class="pill good">reviewed</span>{% endif %}{% else %}<span class="pill warn">off</span>{% endif %}</h2>
  {% if not f.ai.enabled %}
    <p class="muted">{{ f.ai.error }}</p>
  {% else %}
    {% if f.ai.narrative %}<p style="font-size:15px">{{ f.ai.narrative }}</p>{% endif %}
    <p style="margin-top:14px"><strong>FRS observations</strong></p>
    {% if f.ai.frs_observations %}
    <table>
      <thead><tr><th>Severity</th><th>Area</th><th>Issue &amp; detail</th><th>Recommendation</th></tr></thead>
      <tbody>
        {% for o in f.ai.frs_observations %}
        <tr>
          <td>{% set sev = (o.severity or '')|lower %}
            {% if sev == 'high' %}<span class="pill bad">High</span>
            {% elif sev == 'medium' %}<span class="pill warn">Medium</span>
            {% elif sev == 'low' %}<span class="pill">Low</span>{% endif %}</td>
          <td>{{ o.area or o.frs }}{% if o.frs and o.area %}<br><span class="muted">{{ o.frs }}</span>{% endif %}</td>
          <td><strong>{{ o.issue }}</strong>{% if o.detail %}<br><span class="muted">{{ o.detail }}</span>{% endif %}</td>
          <td>{{ o.recommendation }}</td>
        </tr>
        {% endfor %}
      </tbody>
    </table>
    {% else %}<p class="muted">The AI reviewer did not raise FRS observations on this set.</p>{% endif %}
    {% if f.ai.corrected_figures %}
    <p style="margin-top:14px"><strong>Corrected figures</strong> <span class="muted">— ready to drop in</span></p>
    {% for c in f.ai.corrected_figures %}
    <div style="border:1px solid #d7deea;border-radius:8px;padding:12px 14px;margin:8px 0;background:#f7f9fc">
      <div style="font-weight:600;color:#1F3864">{{ c.statement }}</div>
      {% if c.issue %}<div class="muted" style="margin:2px 0 8px">{{ c.issue }}</div>{% endif %}
      <pre style="white-space:pre-wrap;font-family:Consolas,monospace;font-size:13px;margin:0;background:#fff;border:1px solid #e5e9f0;border-radius:6px;padding:10px">{{ c.corrected }}</pre>
      {% if c.note %}<div class="muted" style="margin-top:6px">{{ c.note }}</div>{% endif %}
    </div>
    {% endfor %}
    {% endif %}
    {% if f.ai.suggested_wording %}
    <p style="margin-top:14px"><strong>Suggested replacement wording</strong> <span class="muted">— paste-ready</span></p>
    {% for s in f.ai.suggested_wording %}
    <div style="border-left:4px solid #2E5496;background:#eaf1fb;border-radius:0 8px 8px 0;padding:10px 14px;margin:8px 0">
      <div style="font-weight:600;color:#1F3864;margin-bottom:4px">{{ s.note }}</div>
      <div style="color:#203864;white-space:pre-wrap">{{ s.draft }}</div>
    </div>
    {% endfor %}
    {% endif %}
    <p style="margin-top:14px"><strong>Grammar &amp; drafting</strong></p>
    {% if f.ai.grammar_issues %}
    <table>
      <thead><tr><th>Location</th><th>Current</th><th>Suggested</th></tr></thead>
      <tbody>
        {% for g in f.ai.grammar_issues %}
        <tr><td>{{ g.location }}</td><td>{{ g.current }}</td><td>{{ g.suggested }}</td></tr>
        {% endfor %}
      </tbody>
    </table>
    {% else %}<p class="muted">No drafting issues raised.</p>{% endif %}
  {% endif %}
</div>

{% if f.disclosure_templates %}
<div class="card">
  <h2>Disclosure templates <span class="pill">{{ f.disclosure_templates|length }} triggered</span> <span class="muted" style="font-size:13px;font-weight:400">— paste-ready, based on this review's findings</span></h2>
  <p class="muted">Adapt the [bracketed] items to the company's facts, then paste into the notes. Figures shown were read from this set.</p>
  {% for t in f.disclosure_templates %}
  <div style="border-left:4px solid #2E5496;background:#eaf1fb;border-radius:0 8px 8px 0;padding:10px 14px;margin:10px 0">
    <div style="font-weight:600;color:#1F3864;margin-bottom:4px">{{ t.title }}</div>
    <div style="color:#203864;white-space:pre-wrap;font-size:14px">{{ t.body }}</div>
  </div>
  {% endfor %}
</div>
{% endif %}

<div class="card">
  <h2>Language &amp; spelling (rule scan) {% if f.language_issues %}<span class="pill warn">{{ f.language_issues|length }} to review</span>{% else %}<span class="pill good">Nothing flagged</span>{% endif %}</h2>
  {% if f.language_issues %}
  <table>
    <thead><tr><th>Type</th><th>Found</th><th>Suggest</th><th>Context</th></tr></thead>
    <tbody>
      {% for g in f.language_issues %}
      <tr>
        <td>{{ g.kind }}</td><td>{{ g.found }}</td><td>{{ g.suggest }}</td>
        <td class="muted">…{{ g.context }}…</td>
      </tr>
      {% endfor %}
    </tbody>
  </table>
  {% else %}
  <p class="muted">No British-English or common-typo issues detected by the rule scan.</p>
  {% endif %}
</div>

{% if f.warnings %}
<div class="card">
  <h2>Notes &amp; scope</h2>
  <ul>{% for w in f.warnings %}<li class="muted">{{ w }}</li>{% endfor %}</ul>
</div>
{% endif %}

{% endif %}

<div class="card" style="background:#fffbeb;border-color:#fde68a">
  <p class="muted" style="margin:0">This automated review is a first-pass aid, not a substitute for a full
  FRS/IFRS compliance review by a qualified reviewer.</p>
</div>

<form method="post" action="{{ url_for('delete', rec_id=r.id) }}"
      onsubmit="return confirm('Delete this file and its report?')">
  <button class="btn danger" type="submit">Delete file</button>
</form>
{% endblock %}"""

ADMIN_HTML = """{% extends "base.html" %}
{% block title %}People · FS Review{% endblock %}
{% block content %}
<div class="card">
  <p class="muted"><a href="{{ url_for('dashboard') }}">← Back to dashboard</a></p>
  <h1>People who can sign in</h1>
  <p class="muted">Only people listed here can access the portal. Set a password to give someone a direct email + password login, and share it with them.{% if magic %} Or leave the password blank — they can then request a one-time sign-in link by email.{% endif %}</p>
  <form method="post" action="{{ url_for('admin_invite') }}" style="margin:16px 0">
    <div style="display:flex;gap:10px;flex-wrap:wrap">
      <input type="text" name="name" placeholder="Full name" style="flex:1;min-width:130px">
      <input type="text" name="email" placeholder="email@company.com" required style="flex:1;min-width:180px">
      <input type="text" name="password" placeholder="Set a password" style="flex:1;min-width:140px">
      <button class="btn" type="submit">Add person</button>
    </div>
    <p class="muted" style="margin-top:8px;font-size:13px">Tip: pick a password, click "Add person", then send the person their email + password so they can sign in.</p>
  </form>
  <table>
    <thead><tr><th>Name</th><th>Email / username</th><th>Login</th><th>Can invite</th><th></th></tr></thead>
    <tbody>
      {% for u in users %}
      <tr><td>{{ u.name }}</td><td>{{ u.username }}</td>
        <td>{% if u.has_password %}<span class="pill good">Password</span>{% else %}<span class="pill warn">Email code/link</span>{% endif %}</td>
        <td>{% if u.username == admin_email or u.username == 'admin' %}<span class="pill good">Admin</span>
          {% elif u.can_invite %}<span class="pill good">Yes</span>
            {% if is_full_admin %}<form method="post" action="{{ url_for('admin_grant') }}" style="margin:4px 0 0;display:inline">
              <input type="hidden" name="email" value="{{ u.username }}"><input type="hidden" name="allow" value="0">
              <button class="btn secondary" type="submit" style="padding:2px 8px;font-size:12px">Revoke</button></form>{% endif %}
          {% else %}<span class="muted">No</span>
            {% if is_full_admin %}<form method="post" action="{{ url_for('admin_grant') }}" style="margin:4px 0 0;display:inline">
              <input type="hidden" name="email" value="{{ u.username }}"><input type="hidden" name="allow" value="1">
              <button class="btn secondary" type="submit" style="padding:2px 8px;font-size:12px">Grant</button></form>{% endif %}
          {% endif %}</td>
        <td>{% if is_full_admin and u.username != admin_email and u.username != 'admin' %}
          <form method="post" action="{{ url_for('admin_remove') }}" onsubmit="return confirm('Remove this person?')" style="margin:0">
            <input type="hidden" name="email" value="{{ u.username }}">
            <button class="btn danger" type="submit">Remove</button>
          </form>{% elif u.username == admin_email or u.username == 'admin' %}<span class="muted">admin</span>{% endif %}</td></tr>
      {% endfor %}
    </tbody>
  </table>
  <p class="muted" style="margin-top:10px;font-size:13px">People with "Can invite: Yes" can add new people here, but only the admin can remove people or change invite rights.</p>
</div>
{% endblock %}"""

AUDIT_HTML = """{% extends "base.html" %}
{% block title %}Audit log · FS Review{% endblock %}
{% block content %}
<div class="card">
  <p class="muted"><a href="{{ url_for('dashboard') }}">← Back to dashboard</a></p>
  <h1>Audit log</h1>
  <p class="muted">Every sign-in, upload, review, download, deletion and people change — newest first (latest 500 shown).
     <a class="btn secondary" href="{{ url_for('admin_audit_csv') }}" style="margin-left:10px">Download full log (CSV)</a>
     <a class="btn secondary" href="{{ url_for('admin_backup') }}" style="margin-left:6px">Download backup (JSON)</a></p>
  <table>
    <thead><tr><th>Time (UTC)</th><th>Person</th><th>Action</th><th>Detail</th><th>IP</th></tr></thead>
    <tbody>
      {% for e in events %}
      <tr>
        <td style="white-space:nowrap">{{ e.at }}</td>
        <td>{{ e.actor }}</td>
        <td>{% if 'failed' in e.action %}<span class="pill bad">{{ e.action }}</span>
            {% elif e.action in ('login','logout') %}<span class="pill">{{ e.action }}</span>
            {% elif 'download' in e.action %}<span class="pill good">{{ e.action }}</span>
            {% elif e.action in ('delete record','remove person') %}<span class="pill warn">{{ e.action }}</span>
            {% else %}<span class="pill good">{{ e.action }}</span>{% endif %}</td>
        <td class="muted">{{ e.detail }}</td>
        <td class="muted">{{ e.ip }}</td>
      </tr>
      {% endfor %}
      {% if not events %}<tr><td colspan="5" class="muted">No events logged yet.</td></tr>{% endif %}
    </tbody>
  </table>
</div>
{% endblock %}"""

app.jinja_loader = DictLoader({
    "base.html": BASE_HTML,
    "login.html": LOGIN_HTML,
    "otp.html": OTP_HTML,
    "dashboard.html": DASHBOARD_HTML,
    "report.html": REPORT_HTML,
    "admin.html": ADMIN_HTML,
    "audit.html": AUDIT_HTML,
})


# --------------------------------------------------------------------------
# Tiny JSON "database" helpers
# --------------------------------------------------------------------------
def _load(path, default):
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return default


def _save(path, data):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)


def load_users():
    """Users are { username: {password_hash, name} }.

    On first run we seed an admin account. The default password should be
    changed immediately (or set ADMIN_PASSWORD before first launch).
    """
    users = _load(USERS_FILE, None)
    if users is None:
        admin_pw = os.environ.get("ADMIN_PASSWORD", "ChangeMe123!")
        users = {
            "admin": {
                "password_hash": generate_password_hash(admin_pw),
                "name": "Administrator",
            }
        }
        _save(USERS_FILE, users)
    return users


def load_records():
    return _load(RECORDS_FILE, [])


def save_records(records):
    _save(RECORDS_FILE, records)


# --------------------------------------------------------------------------
# Storage abstraction: Postgres (Neon/Supabase) when DATABASE_URL is set,
# else the JSON-file backend above. Keeps users, reviewed records and the
# original uploaded files so everything persists across restarts/redeploys.
# --------------------------------------------------------------------------
_REC_COLS = ["id", "original_name", "ext", "size_bytes",
             "uploaded_by", "uploaded_at", "findings"]


def _db():
    import psycopg2
    return psycopg2.connect(DATABASE_URL, connect_timeout=10)


def init_db():
    if not USE_DB:
        load_users()           # seed admin in file mode
        return
    conn = _db()
    cur = conn.cursor()
    cur.execute("CREATE TABLE IF NOT EXISTS users ("
                "username TEXT PRIMARY KEY, password_hash TEXT NOT NULL, name TEXT)")
    cur.execute("CREATE TABLE IF NOT EXISTS records ("
                "id TEXT PRIMARY KEY, original_name TEXT, ext TEXT, size_bytes BIGINT, "
                "uploaded_by TEXT, uploaded_at TEXT, findings JSONB, file_bytes BYTEA)")
    cur.execute("CREATE TABLE IF NOT EXISTS login_tokens ("
                "token TEXT PRIMARY KEY, email TEXT, name TEXT, expires_at TEXT)")
    # Invite rights: admins can grant selected people the right to invite others.
    cur.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS can_invite BOOLEAN DEFAULT FALSE")
    # Audit trail: every significant action, kept permanently.
    cur.execute("CREATE TABLE IF NOT EXISTS audit_log ("
                "id SERIAL PRIMARY KEY, at TEXT, actor TEXT, action TEXT, "
                "detail TEXT, ip TEXT)")
    # One-time 6-digit sign-in codes (OTP), one active code per email.
    cur.execute("CREATE TABLE IF NOT EXISTS otp_codes ("
                "email TEXT PRIMARY KEY, code_hash TEXT, expires_at TEXT, "
                "attempts INT DEFAULT 0)")
    cur.execute("SELECT COUNT(*) FROM users")
    if cur.fetchone()[0] == 0:
        pw = os.environ.get("ADMIN_PASSWORD", "ChangeMe123!")
        cur.execute("INSERT INTO users (username, password_hash, name) VALUES (%s,%s,%s)",
                    ("admin", generate_password_hash(pw), "Administrator"))
    # For magic-link login, the allow-list IS the users table (keyed by email).
    # Seed the admin email so the owner can always request a link.
    if ADMIN_EMAIL:
        cur.execute("INSERT INTO users (username, password_hash, name) VALUES (%s,%s,%s) "
                    "ON CONFLICT (username) DO NOTHING",
                    (ADMIN_EMAIL, "", "Administrator"))
    conn.commit()
    cur.close()
    conn.close()


def get_user(username):
    if USE_DB:
        conn = _db()
        cur = conn.cursor()
        cur.execute("SELECT password_hash, name, can_invite FROM users WHERE username=%s",
                    (username,))
        row = cur.fetchone()
        cur.close()
        conn.close()
        return ({"password_hash": row[0], "name": row[1], "can_invite": bool(row[2])}
                if row else None)
    return load_users().get(username)


def add_user(username, password, name):
    if USE_DB:
        conn = _db()
        cur = conn.cursor()
        cur.execute("INSERT INTO users (username, password_hash, name) VALUES (%s,%s,%s) "
                    "ON CONFLICT (username) DO UPDATE SET "
                    "password_hash=EXCLUDED.password_hash, name=EXCLUDED.name",
                    (username, generate_password_hash(password), name))
        conn.commit()
        cur.close()
        conn.close()
    else:
        users = load_users()
        users[username] = {"password_hash": generate_password_hash(password), "name": name}
        _save(USERS_FILE, users)


def invite_user(email, name, password=""):
    """Add an email to the allow-list. If a password is given, the person can
    sign in with email + password right away (no email needed). If left blank,
    they must use the one-time email sign-in link (when that's enabled)."""
    email = email.strip().lower()
    pw_hash = generate_password_hash(password) if password else ""
    if USE_DB:
        conn = _db()
        cur = conn.cursor()
        # Update the name always; only overwrite the password when a new one is given
        # (so re-adding someone without a password doesn't wipe their existing one).
        cur.execute(
            "INSERT INTO users (username, password_hash, name) VALUES (%s,%s,%s) "
            "ON CONFLICT (username) DO UPDATE SET name=EXCLUDED.name, "
            "password_hash = CASE WHEN EXCLUDED.password_hash <> '' "
            "THEN EXCLUDED.password_hash ELSE users.password_hash END",
            (email, pw_hash, name))
        conn.commit()
        cur.close()
        conn.close()
    else:
        users = load_users()
        existing = users.get(email, {})
        users[email] = {"password_hash": pw_hash or existing.get("password_hash", ""),
                        "name": name,
                        "can_invite": bool(existing.get("can_invite"))}
        _save(USERS_FILE, users)


def list_all_users():
    if USE_DB:
        conn = _db()
        cur = conn.cursor()
        cur.execute("SELECT username, name, password_hash, can_invite "
                    "FROM users ORDER BY username")
        rows = cur.fetchall()
        cur.close()
        conn.close()
        return [{"username": r[0], "name": r[1], "has_password": bool(r[2]),
                 "can_invite": bool(r[3])} for r in rows]
    return [{"username": u, "name": v.get("name", ""),
             "has_password": bool(v.get("password_hash")),
             "can_invite": bool(v.get("can_invite"))}
            for u, v in load_users().items()]


def set_invite_right(email, allow):
    """Grant or revoke the right to invite others (admin action)."""
    email = email.strip().lower()
    if USE_DB:
        conn = _db()
        cur = conn.cursor()
        cur.execute("UPDATE users SET can_invite=%s WHERE username=%s", (bool(allow), email))
        conn.commit()
        cur.close()
        conn.close()
    else:
        users = load_users()
        if email in users:
            users[email]["can_invite"] = bool(allow)
            _save(USERS_FILE, users)


def remove_user(email):
    email = email.strip().lower()
    if USE_DB:
        conn = _db()
        cur = conn.cursor()
        cur.execute("DELETE FROM users WHERE username=%s", (email,))
        conn.commit()
        cur.close()
        conn.close()
    else:
        users = load_users()
        users.pop(email, None)
        _save(USERS_FILE, users)


# --------------------------------------------------------------------------
# Audit trail — who did what, when, from where. Never allowed to break the app.
# --------------------------------------------------------------------------
def log_event(action, detail=""):
    try:
        actor = (session.get("user") or "anonymous")
    except Exception:
        actor = "system"
    try:
        ip = (request.headers.get("X-Forwarded-For", "").split(",")[0].strip()
              or request.remote_addr or "")
    except Exception:
        ip = ""
    at = dt.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
    try:
        if USE_DB:
            conn = _db()
            cur = conn.cursor()
            cur.execute("INSERT INTO audit_log (at, actor, action, detail, ip) "
                        "VALUES (%s,%s,%s,%s,%s)",
                        (at, actor, action, str(detail)[:500], ip))
            conn.commit()
            cur.close()
            conn.close()
        else:
            log = _load(AUDIT_FILE, [])
            log.append({"at": at, "actor": actor, "action": action,
                        "detail": str(detail)[:500], "ip": ip})
            _save(AUDIT_FILE, log[-5000:])   # cap file mode at 5,000 entries
    except Exception as e:
        print(f"[audit] could not log '{action}': {e}")


def get_audit_log(limit=500):
    """Most recent audit entries, newest first."""
    if USE_DB:
        conn = _db()
        cur = conn.cursor()
        cur.execute("SELECT at, actor, action, detail, ip FROM audit_log "
                    "ORDER BY id DESC LIMIT %s", (limit,))
        rows = cur.fetchall()
        cur.close()
        conn.close()
        return [{"at": r[0], "actor": r[1], "action": r[2],
                 "detail": r[3], "ip": r[4]} for r in rows]
    log = _load(AUDIT_FILE, [])
    return list(reversed(log[-limit:]))


def is_admin():
    u = (session.get("user") or "").lower()
    return u == "admin" or (ADMIN_EMAIL and u == ADMIN_EMAIL)


def current_can_invite():
    """Admins always can invite; others only if granted the right.

    Cached in the session so this does not cost a database round-trip on every
    single page render (a big win on the free tier, where each new connection
    to the database takes a TLS handshake). Grant/revoke takes effect on the
    person's next sign-in."""
    if is_admin():
        return True
    u = (session.get("user") or "").lower()
    if not u:
        return False
    if "can_invite" not in session:
        user = get_user(u)
        session["can_invite"] = bool(user and user.get("can_invite"))
    return bool(session.get("can_invite"))


def save_record(record, file_bytes):
    if USE_DB:
        import psycopg2
        import psycopg2.extras
        conn = _db()
        cur = conn.cursor()
        cur.execute("INSERT INTO records (id, original_name, ext, size_bytes, uploaded_by, "
                    "uploaded_at, findings, file_bytes) VALUES (%s,%s,%s,%s,%s,%s,%s,%s)",
                    (record["id"], record["original_name"], record["ext"],
                     record["size_bytes"], record["uploaded_by"], record["uploaded_at"],
                     psycopg2.extras.Json(record["findings"]),
                     psycopg2.Binary(file_bytes) if file_bytes else None))
        conn.commit()
        cur.close()
        conn.close()
    else:
        records = load_records()
        records.append(record)
        save_records(records)


def list_records():
    """Dashboard list — metadata only. The findings JSON can be megabytes per
    record; fetching it for every row made the dashboard slow, and the list view
    never displays it."""
    if USE_DB:
        conn = _db()
        cur = conn.cursor()
        cur.execute("SELECT id, original_name, ext, size_bytes, uploaded_by, uploaded_at "
                    "FROM records ORDER BY uploaded_at DESC")
        rows = cur.fetchall()
        cur.close()
        conn.close()
        return [dict(zip(_REC_COLS[:-1], r)) for r in rows]
    return sorted(load_records(), key=lambda r: r["uploaded_at"], reverse=True)


def get_record(rec_id):
    if USE_DB:
        conn = _db()
        cur = conn.cursor()
        cur.execute("SELECT id, original_name, ext, size_bytes, uploaded_by, uploaded_at, "
                    "findings FROM records WHERE id=%s", (rec_id,))
        r = cur.fetchone()
        cur.close()
        conn.close()
        return dict(zip(_REC_COLS, r)) if r else None
    return next((r for r in load_records() if r["id"] == rec_id), None)


def get_record_file(rec_id):
    """Return (original_name, bytes) for download, or None."""
    if USE_DB:
        conn = _db()
        cur = conn.cursor()
        cur.execute("SELECT original_name, file_bytes FROM records WHERE id=%s", (rec_id,))
        r = cur.fetchone()
        cur.close()
        conn.close()
        if r and r[1] is not None:
            return (r[0], bytes(r[1]))
        return None
    rec = get_record(rec_id)
    if rec and rec.get("stored_name"):
        p = os.path.join(UPLOAD_DIR, rec["stored_name"])
        if os.path.exists(p):
            with open(p, "rb") as f:
                return (rec["original_name"], f.read())
    return None


def delete_record(rec_id):
    if USE_DB:
        conn = _db()
        cur = conn.cursor()
        cur.execute("DELETE FROM records WHERE id=%s", (rec_id,))
        conn.commit()
        cur.close()
        conn.close()
    else:
        records = load_records()
        rec = next((r for r in records if r["id"] == rec_id), None)
        if rec:
            try:
                os.remove(os.path.join(UPLOAD_DIR, rec.get("stored_name", "")))
            except OSError:
                pass
            save_records([r for r in records if r["id"] != rec_id])


# --------------------------------------------------------------------------
# Magic-link helpers (token store in DB + email via Resend)
# --------------------------------------------------------------------------
def send_email(to_addr, subject, html):
    """Send an email via the Resend API. Returns True on success."""
    import urllib.request, urllib.error
    if not RESEND_API_KEY:
        return False
    payload = json.dumps({
        "from": MAIL_FROM, "to": [to_addr], "subject": subject, "html": html,
    }).encode("utf-8")
    req = urllib.request.Request(
        "https://api.resend.com/emails", data=payload,
        headers={"Authorization": "Bearer " + RESEND_API_KEY.strip(),
                 "Content-Type": "application/json",
                 # Resend's API is behind Cloudflare, which blocks the default
                 # Python-urllib client signature (Cloudflare error 1010). A normal
                 # browser-style User-Agent gets the request through.
                 "User-Agent": "Mozilla/5.0 (compatible; FS-Review/1.0; +https://fs-review.onrender.com)"})
    try:
        with urllib.request.urlopen(req, timeout=15) as resp:
            return resp.status in (200, 201)
    except urllib.error.HTTPError as e:
        try:
            body = e.read().decode("utf-8", "replace")
        except Exception:
            body = ""
        print(f"[send_email] HTTP {e.code} from={MAIL_FROM} to={to_addr} :: {body[:500]}")
        return False
    except Exception as e:
        print(f"[send_email] {type(e).__name__}: {e}")
        return False


def create_login_token(email, name):
    token = secrets.token_urlsafe(32)
    expires = (dt.datetime.utcnow() + dt.timedelta(minutes=TOKEN_TTL_MIN)).isoformat()
    conn = _db()
    cur = conn.cursor()
    cur.execute("INSERT INTO login_tokens (token, email, name, expires_at) "
                "VALUES (%s,%s,%s,%s)", (token, email, name, expires))
    conn.commit()
    cur.close()
    conn.close()
    return token


def consume_login_token(token):
    """Validate a one-time token; if good, delete it and return (email, name)."""
    conn = _db()
    cur = conn.cursor()
    cur.execute("SELECT email, name, expires_at FROM login_tokens WHERE token=%s", (token,))
    row = cur.fetchone()
    result = None
    if row:
        email, name, expires_at = row
        cur.execute("DELETE FROM login_tokens WHERE token=%s", (token,))
        conn.commit()
        try:
            if dt.datetime.fromisoformat(expires_at) >= dt.datetime.utcnow():
                result = (email, name)
        except Exception:
            result = None
    cur.close()
    conn.close()
    return result


OTP_TTL_MIN = 10       # one-time code validity, minutes
OTP_MAX_ATTEMPTS = 5   # wrong tries before the code is invalidated


def create_otp(email):
    """Generate a 6-digit one-time sign-in code for an invited email (stored hashed)."""
    code = f"{secrets.randbelow(1000000):06d}"
    expires = (dt.datetime.utcnow() + dt.timedelta(minutes=OTP_TTL_MIN)).isoformat()
    conn = _db()
    cur = conn.cursor()
    cur.execute("INSERT INTO otp_codes (email, code_hash, expires_at, attempts) "
                "VALUES (%s,%s,%s,0) "
                "ON CONFLICT (email) DO UPDATE SET code_hash=EXCLUDED.code_hash, "
                "expires_at=EXCLUDED.expires_at, attempts=0",
                (email, generate_password_hash(code), expires))
    conn.commit()
    cur.close()
    conn.close()
    return code


def verify_otp(email, code):
    """Check a 6-digit code. Deletes it on success, expiry or too many attempts."""
    conn = _db()
    cur = conn.cursor()
    cur.execute("SELECT code_hash, expires_at, attempts FROM otp_codes WHERE email=%s",
                (email,))
    row = cur.fetchone()
    ok = False
    if row:
        code_hash, expires_at, attempts = row
        expired = True
        try:
            expired = dt.datetime.fromisoformat(expires_at) < dt.datetime.utcnow()
        except Exception:
            pass
        if expired or attempts >= OTP_MAX_ATTEMPTS:
            cur.execute("DELETE FROM otp_codes WHERE email=%s", (email,))
        elif check_password_hash(code_hash, code):
            ok = True
            cur.execute("DELETE FROM otp_codes WHERE email=%s", (email,))
        else:
            cur.execute("UPDATE otp_codes SET attempts=attempts+1 WHERE email=%s", (email,))
        conn.commit()
    cur.close()
    conn.close()
    return ok


# Simple in-memory throttle: cap magic-link requests per email (per worker).
_link_requests = {}


def too_many_requests(key, limit=5, window_min=15):
    now = dt.datetime.utcnow()
    hits = [t for t in _link_requests.get(key, []) if (now - t).total_seconds() < window_min * 60]
    hits.append(now)
    _link_requests[key] = hits
    return len(hits) > limit


# --------------------------------------------------------------------------
# Auth
# --------------------------------------------------------------------------
def login_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if "user" not in session:
            return redirect(url_for("login", next=request.path))
        return view(*args, **kwargs)
    return wrapped


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        # The identifier may come from either the email field or the legacy
        # username field. Password is optional (blank = request an email link).
        ident = (request.form.get("email") or request.form.get("username") or "").strip().lower()
        password = request.form.get("password", "")
        send_link = request.form.get("sendlink")
        send_otp = request.form.get("sendotp")

        # 0) One-time 6-digit code (OTP) by email — the recommended sign-in.
        if MAGIC_LOGIN and send_otp and ident:
            if too_many_requests(ident):
                flash("Too many requests — please wait a few minutes and try again.", "error")
                return render_template("login.html", magic=True)
            user = get_user(ident)            # the allow-list check
            if user:
                code = create_otp(ident)
                send_email(
                    ident, "Your FS Review sign-in code",
                    f"<p>Hello,</p><p>Your one-time sign-in code is:</p>"
                    f"<p style='font-size:28px;font-weight:bold;letter-spacing:6px'>{code}</p>"
                    f"<p>It expires in {OTP_TTL_MIN} minutes. If you didn't request it, "
                    f"you can ignore this email.</p>")
            # Same message either way — don't reveal whether the email is invited.
            flash("If your email is on the invite list, a 6-digit code is on its way. "
                  "Enter it below.", "success")
            return render_template("otp.html", email=ident)

        # 1) Password sign-in — works for anyone whose account has a password set.
        if password and not send_link:
            user = get_user(ident)
            if user and user.get("password_hash") and \
                    check_password_hash(user["password_hash"], password):
                session["user"] = ident
                session["name"] = user.get("name", ident)
                session["can_invite"] = bool(user.get("can_invite"))
                log_event("login", f"password sign-in: {ident}")
                return redirect(request.args.get("next") or url_for("dashboard"))
            log_event("login failed", f"wrong email/password for: {ident}")
            flash("Invalid email or password.", "error")
            return render_template("login.html", magic=MAGIC_LOGIN)

        # 2) One-time email sign-in link (only when magic-link is configured).
        if MAGIC_LOGIN and ident:
            if too_many_requests(ident):
                flash("Too many requests — please wait a few minutes and try again.", "error")
                return render_template("login.html", magic=True)
            user = get_user(ident)            # the allow-list check
            if user:
                token = create_login_token(ident, user.get("name") or ident)
                base = (os.environ.get("APP_BASE_URL") or request.host_url).rstrip("/")
                if base.startswith("http://"):
                    base = "https://" + base[len("http://"):]
                link = base + url_for("auth_token", token=token)
                send_email(
                    ident, "Your FS Review sign-in link",
                    f"<p>Hello,</p><p>Click below to sign in to the FS Review portal:</p>"
                    f'<p><a href="{link}">Sign in to FS Review</a></p>'
                    f"<p>This link expires in {TOKEN_TTL_MIN} minutes. "
                    f"If you didn't request it, you can ignore this email.</p>")
            # Same message either way — don't reveal whether the email is invited.
            flash("If your email is on the invite list, a sign-in link is on its way. "
                  "Check your inbox.", "success")
            return render_template("login.html", magic=MAGIC_LOGIN)

        flash("Enter your email and password to sign in.", "error")
    return render_template("login.html", magic=MAGIC_LOGIN)


@app.route("/otp", methods=["POST"])
def otp_verify():
    """Second step of the OTP sign-in: check the 6-digit code."""
    email = (request.form.get("email") or "").strip().lower()
    code = (request.form.get("code") or "").strip().replace(" ", "")
    if not (MAGIC_LOGIN and email and code):
        flash("Enter the 6-digit code from your email.", "error")
        return render_template("otp.html", email=email)
    if verify_otp(email, code):
        user = get_user(email)
        session["user"] = email
        session["name"] = (user or {}).get("name") or email
        session["can_invite"] = bool((user or {}).get("can_invite"))
        log_event("login", f"OTP sign-in: {email}")
        return redirect(url_for("dashboard"))
    log_event("login failed", f"wrong/expired OTP for: {email}")
    flash("That code is incorrect or has expired — request a new one.", "error")
    return render_template("otp.html", email=email)


@app.route("/auth/<token>")
def auth_token(token):
    got = consume_login_token(token) if MAGIC_LOGIN else None
    if got:
        email, name = got
        session["user"] = email
        session["name"] = name or email
        log_event("login", f"email-link sign-in: {email}")
        return redirect(url_for("dashboard"))
    flash("That sign-in link is invalid or has expired — please request a new one.", "error")
    return redirect(url_for("login"))


@app.route("/admin")
@login_required
def admin_users():
    if not current_can_invite():
        abort(403)
    return render_template("admin.html", users=list_all_users(),
                           magic=MAGIC_LOGIN, admin_email=ADMIN_EMAIL,
                           is_full_admin=is_admin())


@app.route("/admin/invite", methods=["POST"])
@login_required
def admin_invite():
    if not current_can_invite():
        abort(403)
    email = request.form.get("email", "").strip().lower()
    name = request.form.get("name", "").strip() or email
    password = request.form.get("password", "").strip()
    if email:
        invite_user(email, name, password)
        log_event("invite", f"added {email}" + (" (with password)" if password else " (email sign-in)"))
        if password:
            flash(f"Added {email}. They can sign in now with that email and password — "
                  f"send those details to them.", "success")
        else:
            flash(f"Added {email}. They can request a one-time email sign-in link.", "success")
    return redirect(url_for("admin_users"))


@app.route("/admin/remove", methods=["POST"])
@login_required
def admin_remove():
    if not is_admin():
        abort(403)
    email = request.form.get("email", "").strip().lower()
    if email and email != (ADMIN_EMAIL or "") and email != "admin":
        remove_user(email)
        log_event("remove person", email)
        flash(f"Removed {email}.", "success")
    return redirect(url_for("admin_users"))


@app.route("/admin/grant", methods=["POST"])
@login_required
def admin_grant():
    """Grant or revoke invite rights (full admins only)."""
    if not is_admin():
        abort(403)
    email = request.form.get("email", "").strip().lower()
    allow = request.form.get("allow") == "1"
    if email and email != (ADMIN_EMAIL or "") and email != "admin":
        set_invite_right(email, allow)
        log_event("invite rights", ("granted to " if allow else "revoked from ") + email)
        flash(("Granted invite rights to " if allow else "Revoked invite rights from ")
              + email + ".", "success")
    return redirect(url_for("admin_users"))


@app.route("/admin/audit")
@login_required
def admin_audit():
    """Audit trail — full admin only."""
    if not is_admin():
        abort(403)
    return render_template("audit.html", events=get_audit_log(500))


@app.route("/admin/audit.csv")
@login_required
def admin_audit_csv():
    if not is_admin():
        abort(403)
    import csv as _csv
    from flask import Response
    buf = io.StringIO()
    w = _csv.writer(buf)
    w.writerow(["Time (UTC)", "Person", "Action", "Detail", "IP"])
    for e in get_audit_log(5000):
        w.writerow([e["at"], e["actor"], e["action"], e["detail"], e["ip"]])
    return Response(buf.getvalue(), mimetype="text/csv",
                    headers={"Content-Disposition":
                             "attachment; filename=fs-review-audit-log.csv"})


@app.errorhandler(500)
def server_error(e):
    """Friendly error page + the failure goes into the audit log so problems
    are visible instead of silent."""
    try:
        log_event("error", f"500 on {request.path}: {e}")
    except Exception:
        pass
    return ("<div style='font-family:Arial;max-width:560px;margin:80px auto;"
            "text-align:center'><h2>Something went wrong</h2>"
            "<p>The error has been logged. Please try again — if it keeps "
            "happening, tell your administrator what you were doing.</p>"
            "<p><a href='/'>Back to the dashboard</a></p></div>"), 500


@app.errorhandler(413)
def too_large(e):
    flash("That file is larger than the 25 MB limit.", "error")
    return redirect(url_for("dashboard"))


@app.route("/admin/backup.json")
@login_required
def admin_backup():
    """One-click backup (admin only): people, records metadata + findings, and
    the audit trail as a single JSON file. The uploaded source files themselves
    live in the database (Neon), which is the primary store."""
    if not is_admin():
        abort(403)
    from flask import Response
    users = [{k: v for k, v in u.items()} for u in list_all_users()]
    records = list_records()
    payload = {
        "exported_at": dt.datetime.utcnow().isoformat() + "Z",
        "users": users,
        "records": records,
        "audit_log": get_audit_log(5000),
    }
    log_event("backup export")
    return Response(json.dumps(payload, indent=2, default=str),
                    mimetype="application/json",
                    headers={"Content-Disposition":
                             "attachment; filename=fs-review-backup.json"})


@app.route("/healthz")
def healthz():
    """Lightweight keep-alive endpoint (no database hit) — point a free uptime
    pinger at this every ~10 minutes to stop the free instance spinning down."""
    return "ok", 200


# Firm branding: drop a "logo.png" next to app.py (commit it to the repo) and it
# appears in the portal header and on the Word review report automatically.
LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logo.png")


@app.route("/logo.png")
def logo():
    from flask import send_file
    if not os.path.exists(LOGO_PATH):
        abort(404)
    return send_file(LOGO_PATH, mimetype="image/png")


@app.context_processor
def inject_globals():
    return {"current_is_admin": is_admin(),
            "current_user_can_invite": current_can_invite(),
            "has_logo": os.path.exists(LOGO_PATH)}


@app.after_request
def set_security_headers(resp):
    resp.headers["X-Frame-Options"] = "DENY"
    resp.headers["X-Content-Type-Options"] = "nosniff"
    resp.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    resp.headers["Content-Security-Policy"] = (
        "default-src 'self' 'unsafe-inline'; img-src 'self' data:; "
        "frame-ancestors 'none'; base-uri 'self'; form-action 'self'")
    resp.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
    # Privacy: keep the whole portal out of search engines / AI crawlers.
    resp.headers["X-Robots-Tag"] = "noindex, nofollow, noarchive, nosnippet"
    # Lock down powerful browser features the app never uses.
    resp.headers["Permissions-Policy"] = (
        "camera=(), microphone=(), geolocation=(), interest-cohort=()")
    # Confidential financials must never be cached on shared/public computers
    # or by intermediary proxies. Static assets can still be cached.
    if request.path != "/robots.txt":
        resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, private"
        resp.headers["Pragma"] = "no-cache"
    return resp


@app.route("/robots.txt")
def robots_txt():
    # Tell every crawler to stay out of the entire site.
    return Response("User-agent: *\nDisallow: /\n", mimetype="text/plain")


@app.route("/logout")
def logout():
    log_event("logout")
    session.clear()
    return redirect(url_for("login"))


# --------------------------------------------------------------------------
# Lightweight automated review of a .docx financial statement
# --------------------------------------------------------------------------
KEY_SECTIONS = [
    "statement of financial position",
    "statement of comprehensive income",
    "statement of profit or loss",
    "statement of changes in equity",
    "statement of cash flows",
    "notes to the financial statements",
    "directors' statement",
    "independent auditor",
]

# American -> British spelling (Singapore entities use British English)
US_TO_UK = {
    "organization": "organisation", "organizations": "organisations",
    "recognize": "recognise", "recognized": "recognised", "recognizes": "recognises",
    "recognizing": "recognising", "capitalize": "capitalise",
    "capitalized": "capitalised", "realize": "realise", "realized": "realised",
    "amortize": "amortise", "amortized": "amortised", "utilize": "utilise",
    "utilized": "utilised", "labor": "labour", "color": "colour",
    "center": "centre", "fulfill": "fulfil", "favor": "favour",
    "analyze": "analyse", "analyzed": "analysed", "offset off": "offset",
}

# Common typos seen in financial statements
COMMON_TYPOS = {
    "theses": "these", "finanical": "financial", "financal": "financial",
    "statment": "statement", "statments": "statements",
    "comparitive": "comparative", "accomodate": "accommodate",
    "recieve": "receive", "recievable": "receivable", "seperate": "separate",
    "occured": "occurred", "non-curent": "non-current", "balacne": "balance",
    "liabilites": "liabilities", "expences": "expenses", "incured": "incurred",
    "yearically": "periodically", "theses": "these", "managment": "management",
    "acquisiton": "acquisition", "transalated": "translated",
}

# Incorrect PHRASING commonly seen in SME drafts — corrected as tracked changes
# in the revised FS (case-sensitive so the replacement keeps the sentence intact).
INCORRECT_PHRASES = {
    "The Company are": "The Company is",
    "the Company are": "the Company is",
    "theses standard": "these standards",
    "these standard did": "these standards did",
    "de recognition": "derecognition",
    "cash generating unit": "cash-generating unit",
    "financial statement of the Company": "financial statements of the Company",
    "The financial statement for the financial year":
        "The financial statements for the financial year",
    "board of Directors of the company": "Board of Directors of the Company",
    "adoption of theses standard": "adoption of these standards",
    "material effect on the financial statement of":
        "material effect on the financial statements of",
}

# FRS disclosure checklist — keyword presence heuristics
FRS_CHECKS = [
    ("FRS 1", "Going concern basis stated",
     ["going concern"]),
    ("FRS 1", "Significant judgements & estimates note",
     ["significant judgement", "significant estimate", "judgements and estimates",
      "key sources of estimation"]),
    ("FRS 2", "Inventory cost formula (FIFO / weighted average)",
     ["first-in", "first in", "weighted average", "fifo"]),
    ("FRS 12", "Deferred tax disclosure",
     ["deferred tax"]),
    ("FRS 109", "Financial instruments note",
     ["financial instrument", "financial asset", "financial liabilit"]),
    ("FRS 115", "Revenue recognition timing (over time / point in time)",
     ["over time", "point in time", "performance obligation"]),
    ("FRS 116", "Leases (right-of-use / lease liability)",
     ["right-of-use", "right of use", "lease liabilit", "lease liability"]),
]


def _to_number(text):
    """Parse an accounting-style number; parentheses mean negative."""
    if text is None:
        return None
    t = text.strip().replace("–", "-").replace("−", "-")
    if t in ("", "-", "–", "—", "nil", "Nil", "NIL"):
        return None
    neg = t.startswith("(") and t.endswith(")")
    t = t.strip("()").replace(",", "").replace("$", "").strip()
    try:
        val = float(t)
        return -val if neg else val
    except ValueError:
        return None


def _grid(table):
    """Return (labels, numgrid): first-column text labels and a numeric grid."""
    labels, numgrid = [], []
    for row in table.rows:
        cells = [c.text for c in row.cells]
        labels.append(cells[0].strip() if cells else "")
        numgrid.append([_to_number(c) for c in cells])
    return labels, numgrid


# Grand totals sum *subtotals* across sections, not leaf lines — so they are
# not checked by the simple "sum the lines above" rule (it would misfire).
GRAND_TOTALS = {
    "total assets", "total liabilities", "total equity and liabilities",
    "total equity & liabilities", "net assets", "net current assets",
    "total comprehensive income", "total comprehensive loss",
}


def _note_columns(table):
    """Column indexes that are 'Note' reference columns (values like 7, 8, 9).

    Scans the first several rows for a 'Note' header cell. We must NOT stop at the
    first row containing a year/$, because some statements put the year row ABOVE
    the 'Note' label row — stopping early missed the note column and made downstream
    extraction (cash, share capital) pick up the note number instead of the figure.
    """
    skip = set()
    for row in list(table.rows)[:6]:
        for i, cell in enumerate(row.cells):
            if cell.text.strip().lower() == "note":
                skip.add(i)
    return skip


def check_table_totals(t_idx, table):
    """Check each subtotal against the line items directly above it.

    Conservative on purpose, to avoid false positives on real statements:
    - resets at blank rows and section headers (so balance-sheet sections
      aren't summed together);
    - never treats a 'Total' row as a line item;
    - skips 'Note' reference columns and grand totals;
    - treats a nil dash ('-') as zero, not a section break;
    - flags a sign/brackets inconsistency when the sum equals the stated
      figure in magnitude but the signs differ (e.g. 430,872 vs (430,872)).
    """
    issues = []
    labels, numgrid = _grid(table)
    if len(numgrid) < 3:
        return issues
    ncols = max((len(r) for r in numgrid), default=0)
    skipcols = _note_columns(table)
    moneycols = [c for c in range(1, ncols) if c not in skipcols]
    # A "data row" has a number in at least one money column; otherwise it is a
    # section header / separator.
    is_data = [
        any(numgrid[r][c] is not None for c in moneycols if c < len(numgrid[r]))
        for r in range(len(numgrid))
    ]
    for c in moneycols:
        seg = []
        for r, label in enumerate(labels):
            low = label.lower().strip()
            val = numgrid[r][c] if c < len(numgrid[r]) else None
            if not low:                       # blank row -> section break
                seg = []
                continue
            if "total" in low:                # subtotal / total line
                if low not in GRAND_TOTALS and val is not None and len(seg) >= 2:
                    s = sum(seg)
                    diff = s - val
                    if abs(diff) > 0.5 and val != 0:
                        rel = abs(diff) / (abs(val) + 1e-9)
                        if abs(s + val) <= 0.5:        # same size, opposite sign
                            issues.append({
                                "table": t_idx + 1, "label": label[:55],
                                "sum_of_parts": round(s, 2),
                                "stated_total": round(val, 2),
                                "difference": round(diff, 2),
                                "kind": "sign / brackets",
                            })
                        elif rel < 1.0:                # conservative arithmetic slip
                            issues.append({
                                "table": t_idx + 1, "label": label[:55],
                                "sum_of_parts": round(s, 2),
                                "stated_total": round(val, 2),
                                "difference": round(diff, 2),
                                "kind": "sum",
                            })
                seg = []
                continue                      # never treat a total as a line item
            if not is_data[r]:                # section header with no numbers
                seg = []
                continue
            seg.append(val if val is not None else 0.0)  # nil '-' counts as 0
    return issues


def check_balance_equation(doc):
    """Try to confirm Total assets == Total equity + Total liabilities."""
    results = []
    # Collect labelled totals from all tables, per numeric column
    for table in doc.tables:
        labels, numgrid = _grid(table)
        ncols = max((len(r) for r in numgrid), default=0)
        for c in range(1, ncols):
            picks = {"assets": None, "equity": None, "liabilities": None,
                     "eq_and_liab": None}
            for r, label in enumerate(labels):
                low = label.lower()
                val = numgrid[r][c] if c < len(numgrid[r]) else None
                if val is None:
                    continue
                if "total equity and" in low or "total equity & " in low:
                    picks["eq_and_liab"] = val
                elif "total asset" in low:
                    picks["assets"] = val
                elif "total equit" in low:
                    picks["equity"] = val
                elif "total liabilit" in low:
                    picks["liabilities"] = val
            if picks["assets"] is not None:
                target = picks["eq_and_liab"]
                if target is None and picks["equity"] is not None and picks["liabilities"] is not None:
                    target = picks["equity"] + picks["liabilities"]
                if target is not None:
                    diff = picks["assets"] - target
                    results.append({
                        "total_assets": round(picks["assets"], 2),
                        "equity_plus_liabilities": round(target, 2),
                        "difference": round(diff, 2),
                        "balanced": abs(diff) <= 0.5,
                    })
    # De-duplicate identical results
    seen, uniq = set(), []
    for r in results:
        key = (r["total_assets"], r["equity_plus_liabilities"])
        if key not in seen:
            seen.add(key)
            uniq.append(r)
    return uniq


def _find_row(labels, numgrid, c, *keys, exclude=()):
    """First numeric value in column c whose label matches any key."""
    for r, label in enumerate(labels):
        low = label.lower()
        if any(k in low for k in keys) and not any(x in low for x in exclude):
            v = numgrid[r][c] if c < len(numgrid[r]) else None
            if v is not None:
                return v
    return None


def check_pl(t_idx, table):
    """Profit & loss flow checks (only if the table looks like a P&L)."""
    out = []
    labels, numgrid = _grid(table)
    low_all = " ".join(labels).lower()
    if "gross profit" not in low_all and "gross loss" not in low_all:
        return out
    skip = _note_columns(table)
    ncols = max((len(r) for r in numgrid), default=0)
    for c in [x for x in range(1, ncols) if x not in skip]:
        rev = _find_row(labels, numgrid, c, "revenue", "sales", "turnover", exclude=("cost",))
        cogs = _find_row(labels, numgrid, c, "cost of goods", "cost of sales")
        gp = _find_row(labels, numgrid, c, "gross profit", "gross loss")
        pbt = _find_row(labels, numgrid, c, "before tax", "before taxation")
        tax = _find_row(labels, numgrid, c, "income tax", "tax expense", "taxation")
        net = _find_row(labels, numgrid, c, "for the financial year", "for the year",
                        "loss for", "profit for")
        if rev is not None and cogs is not None and gp is not None:
            exp = rev + cogs if cogs < 0 else rev - cogs
            if abs(exp - gp) > 0.5:
                out.append({"table": t_idx + 1,
                            "check": "Gross profit = Revenue − Cost of sales",
                            "expected": round(exp, 2), "stated": round(gp, 2),
                            "difference": round(exp - gp, 2)})
        # FULL CAST: gross profit + every income/expense line between GP and profit
        # before tax MUST equal profit before tax. Sum the lines strictly between the
        # two, skipping subtotal ("total") and "gross" lines so we don't double-count.
        if gp is not None and pbt is not None:
            gi = pi = None
            for r, label in enumerate(labels):
                low = label.lower()
                if gi is None and ("gross profit" in low or "gross loss" in low):
                    gi = r
                if "before tax" in low or "before taxation" in low:
                    pi = r
            if gi is not None and pi is not None and pi > gi:
                seg = 0.0
                for r in range(gi + 1, pi):
                    low = labels[r].lower()
                    if "total" in low or "gross" in low:
                        continue
                    v = numgrid[r][c] if c < len(numgrid[r]) else None
                    if v is not None:
                        seg += v
                calc = gp + seg
                # Only flag a plausible CASTING SLIP (rounding / transposition): a small
                # difference. A large gap means the parser mis-read a messy layout (note
                # columns, merged/duplicate rows) rather than a real error, so stay silent
                # to avoid false positives — big misstatements are caught by other checks.
                diff = calc - pbt
                if 0.5 < abs(diff) <= 1000:
                    out.append({"table": t_idx + 1,
                                "check": "Profit before tax does not cast (gross profit + "
                                         "income − expenses)",
                                "expected": round(calc, 2), "stated": round(pbt, 2),
                                "difference": round(diff, 2)})
        if net is not None and pbt is not None:
            tx = tax or 0
            # Tax may be shown positive ("1,892") or bracketed ("(1,892)" → parsed
            # negative). Use the magnitude so both presentation styles are handled:
            exp_minus = pbt - abs(tx)     # tax treated as an expense (deduct)
            exp_plus = pbt + abs(tx)      # tax treated as a credit (add back)
            if abs(exp_minus - net) <= 0.5:
                pass                      # ties correctly, nothing to flag
            elif tx and abs(exp_plus - net) <= 0.5:
                # Net = before tax + tax. Arithmetically self-consistent but the tax
                # figure is shown as a positive "expense" while being added back —
                # usually a prior-year over-provision credit shown without brackets.
                # Flag softly as a presentation point, not a hard casting error.
                out.append({"table": t_idx + 1,
                            "check": ("Income tax of {:,.2f} is added back (a net credit), "
                                      "not deducted — acceptable if it is a prior-year "
                                      "over-provision / refund, but it should then be shown "
                                      "in brackets and disclosed as an over-provision line "
                                      "in the tax note; if it is a genuine expense the "
                                      "profit/loss for the year should be "
                                      "{:,.2f}".format(abs(tx), exp_minus)),
                            "expected": round(exp_minus, 2), "stated": round(net, 2),
                            "difference": round(exp_minus - net, 2),
                            "kind": "tax sign / presentation"})
            else:
                out.append({"table": t_idx + 1,
                            "check": "Loss/profit for year = Before tax − Tax",
                            "expected": round(exp_minus, 2), "stated": round(net, 2),
                            "difference": round(exp_minus - net, 2)})
    return out


def check_row_totals(t_idx, table):
    """Horizontal check: when the LAST column is a 'Total' column, each row's
    components should add across to it (e.g. Statement of Changes in Equity).

    Merged cells are counted ONCE: python-docx repeats the same underlying cell
    for each column a merge spans, which used to double-count values and produce
    false 'does not cast across' flags."""
    out = []
    last_is_total = False
    for row in table.rows:
        cells = [c.text.strip().lower() for c in row.cells]
        if cells and "total" in cells[-1]:
            last_is_total = True
            break
    if not last_is_total:
        return out
    for r_idx, row in enumerate(table.rows):
        seen, vals, label = set(), [], ""
        for i, cell in enumerate(row.cells):
            tc = id(cell._tc)
            if tc in seen:
                continue
            seen.add(tc)
            if i == 0 or not label:
                label = label or cell.text.strip()
            v = _to_number(cell.text)
            if i > 0 and v is not None:
                vals.append(v)
        if len(vals) >= 3:                    # at least 2 components + a total
            comps, tot = vals[:-1], vals[-1]
            if abs(sum(comps) - tot) > 0.5:
                out.append({"table": t_idx + 1,
                            "row": (label[:40] or f"row {r_idx + 1}"),
                            "sum_across": round(sum(comps), 2),
                            "stated_total": round(tot, 2),
                            "difference": round(sum(comps) - tot, 2)})
    return out


def check_language(doc):
    """British-English and common-typo checks across all paragraphs."""
    issues = []
    seen = set()
    for p in doc.paragraphs:
        text = p.text
        low = text.lower()
        for us, uk in US_TO_UK.items():
            if re.search(r"\b" + re.escape(us) + r"\b", low) and us not in seen:
                seen.add(us)
                issues.append({"kind": "Spelling (US→UK)",
                               "found": us, "suggest": uk,
                               "context": text.strip()[:90]})
        for phrase, fix in INCORRECT_PHRASES.items():
            if phrase in text and phrase not in seen:
                seen.add(phrase)
                issues.append({"kind": "Incorrect phrasing",
                               "found": phrase, "suggest": fix,
                               "context": text.strip()[:90]})
        for typo, fix in COMMON_TYPOS.items():
            if re.search(r"\b" + re.escape(typo) + r"\b", low) and typo not in seen:
                seen.add(typo)
                issues.append({"kind": "Likely typo",
                               "found": typo, "suggest": fix,
                               "context": text.strip()[:90]})
        if re.search(r"\bSec\.?\s+\d", text) and "Sec." not in seen:
            seen.add("Sec.")
            issues.append({"kind": "Abbreviation",
                           "found": "Sec.", "suggest": "Section",
                           "context": text.strip()[:90]})
    return issues


def check_frs(full_text_low, has_inventory):
    out = []
    for frs, item, keywords in FRS_CHECKS:
        if frs == "FRS 2" and not has_inventory:
            continue
        present = any(k in full_text_low for k in keywords)
        out.append({"frs": frs, "item": item, "present": present})
    return out


# Going-concern disclosure template — the elements a robust FRS 1 going-concern
# note should contain, especially where there are losses / net current liabilities.
GC_ELEMENTS = [
    ("Going-concern basis stated",
     ["going concern basis", "prepared on a going concern", "going concern assumption",
      "basis of going concern"]),
    ("Conditions / indicators disclosed",
     ["net loss", "negative operating cash", "accumulated losses", "net current liabilit",
      "incurred a loss", "net loss of", "operating cash outflow"]),
    ("Material uncertainty acknowledged",
     ["material uncertainty"]),
    ("12-month assessment period stated",
     ["twelve months", "12 months", "next twelve months", "at least twelve months"]),
    ("Management's plans / mitigation",
     ["cost containment", "cost-containment", "shareholder support", "continued financial support",
      "continuing financial support", "additional financing", "raise additional capital",
      "secure additional funding", "obtained an undertaking", "letter of support"]),
    ("Conclusion that the basis is appropriate",
     ["consider this basis to be appropriate", "appropriate", "able to pay its debts",
      "continue as a going concern", "able to continue"]),
]


def check_going_concern(full_text_low):
    """Template check for the going-concern disclosure."""
    mentions_gc = "going concern" in full_text_low
    has_losses = any(k in full_text_low for k in (
        "accumulated losses", "net loss", "net current liabilit",
        "incurred a loss", "negative operating cash", "operating cash outflow"))
    elements = [{"element": label, "present": any(k in full_text_low for k in kws)}
                for label, kws in GC_ELEMENTS]
    return {"mentions_gc": mentions_gc, "has_losses": has_losses, "elements": elements}


# Phrases that show the going-concern basis is supported by external financial
# support — what FRS 1 expects when a company is insolvent / has net current
# liabilities and still prepares accounts on a going-concern basis.
# Kept deliberately SPECIFIC: these phrases evidence that a named party has
# undertaken to support the company. Generic solvency phrases such as "as and
# when they fall due" are NOT here — they are boilerplate in the directors'
# statement and previously produced a false "financial support specified: Yes".
FIN_SUPPORT_KWS = [
    "continued financial support", "continuing financial support",
    "continue to provide financial support", "provide financial support",
    "provide continuing financial support", "financial support from",
    "shareholder support", "shareholders' support", "shareholder's support",
    "will not recall", "will not demand repayment", "not to recall",
    "undertaking to provide", "undertaken to provide", "undertaking not to",
    "letter of support", "letter of financial support",
    "support from its holding", "support from the holding",
    "support from its shareholder", "support from the shareholder",
    "support from its director", "support from the director",
]


def _scan_equity(doc):
    """Find total equity / net assets anywhere in the accounts (any table, any
    money column), so balance-sheet insolvency is caught even when equity sits
    apart from the current-asset lines."""
    for table in doc.tables:
        labels, numgrid = _grid(table)
        skip = _note_columns(table)
        ncols = max((len(r) for r in numgrid), default=0)
        for c in range(1, ncols):
            if c in skip:
                continue
            v = _find_row(labels, numgrid, c,
                          "total equity", "net assets", "net liabilities",
                          "shareholders' equity", "shareholder's equity",
                          exclude=("and liab", "& liab", "and liabilities"))
            if v is not None:
                return v
    return None


def extract_solvency(doc):
    """Pull solvency figures from the statement of financial position: total
    equity (net assets), total current assets and total current liabilities,
    taken from the column where the most could be located (the current year).
    Equity is also scanned independently as a fallback so balance-sheet
    insolvency is caught even if equity sits in another column/table."""
    best = None  # (score, dict)
    for table in doc.tables:
        labels, numgrid = _grid(table)
        low_all = " ".join(labels).lower()
        if ("total equit" not in low_all and "net asset" not in low_all
                and "current liabilit" not in low_all):
            continue
        skip = _note_columns(table)
        ncols = max((len(r) for r in numgrid), default=0)
        for c in range(1, ncols):
            if c in skip:
                continue
            equity = _find_row(labels, numgrid, c,
                               "total equity", "net assets", "net liabilities",
                               "shareholders' equity", "shareholder's equity",
                               exclude=("and liab", "& liab", "and liabilities"))
            ca = _find_row(labels, numgrid, c, "total current assets")
            cl = _find_row(labels, numgrid, c, "total current liabilities")
            ncd = _find_row(labels, numgrid, c,
                            "net current asset", "net current liabilit")
            cand = {"equity": equity, "current_assets": ca,
                    "current_liabilities": cl, "net_current_direct": ncd}
            score = sum(v is not None for v in cand.values())
            if score and (best is None or score > best[0]):
                best = (score, cand)
    out = best[1] if best else {"equity": None, "current_assets": None,
                                "current_liabilities": None,
                                "net_current_direct": None}
    if out["equity"] is None:
        out["equity"] = _scan_equity(doc)
    return out


def assess_going_concern(doc, full_text_low):
    """Going-concern review that tallies the disclosure back to the accounts.

    Solvent (positive equity + net current assets, no losses) -> a standard
    going-concern basis is fine. If the company is in a net-liabilities or
    net-current-liabilities position, or is loss-making, the note must specify
    the financial support it relies on (from shareholders / directors / holding
    company), the material uncertainty, a 12-month assessment and a conclusion."""
    base = check_going_concern(full_text_low)
    sol = extract_solvency(doc)
    equity = sol["equity"]
    ca, cl = sol["current_assets"], sol["current_liabilities"]
    net_current = None
    if ca is not None and cl is not None:
        net_current = round(ca - cl, 2)
    elif sol["net_current_direct"] is not None:
        net_current = round(sol["net_current_direct"], 2)

    bs_insolvent = equity is not None and equity < 0
    liquidity_concern = net_current is not None and net_current < 0
    has_losses = base["has_losses"]
    at_risk = bool(bs_insolvent or liquidity_concern or has_losses)
    fin_support = any(k in full_text_low for k in FIN_SUPPORT_KWS)

    reasons = []
    if bs_insolvent:
        reasons.append("a net liabilities position (total equity {:,.2f})".format(equity))
    if liquidity_concern:
        reasons.append("net current liabilities of {:,.2f}".format(abs(net_current)))
    if has_losses and not (bs_insolvent or liquidity_concern):
        reasons.append("loss / negative operating cash-flow indicators")

    if not at_risk:
        bits = []
        if equity is not None:
            bits.append("net assets are positive ({:,.2f})".format(equity))
        if net_current is not None and net_current >= 0:
            bits.append("net current assets are positive ({:,.2f})".format(net_current))
        lead = (" and ".join(bits) + "; ") if bits else ""
        verdict = (lead + "no loss or insolvency indicators were detected. The company "
                   "appears solvent, so a standard going-concern basis is appropriate "
                   "and no material-uncertainty disclosure is required. Confirm manually.")
        verdict_level = "good"
    else:
        verdict = ("Going-concern RISK — " + "; ".join(reasons) + ". On a going-concern "
                   "basis the note must specify the financial support relied on (e.g. "
                   "continued support from shareholders / directors / the holding company "
                   "and an undertaking not to recall amounts due), acknowledge the material "
                   "uncertainty, state a 12-month assessment period, and give the directors' "
                   "conclusion that the basis remains appropriate.")
        verdict_level = "bad" if (bs_insolvent or liquidity_concern) else "warn"

    contradictions = []
    if at_risk and not base["mentions_gc"]:
        contradictions.append(
            "Going-concern risk indicators are present in the figures but there is no "
            "'going concern' discussion in the notes — this must be addressed.")
    if at_risk and not fin_support:
        contradictions.append(
            "The accounts show a going-concern risk but no statement of financial support "
            "(from shareholders / directors / the holding company, or a letter of support) "
            "was found — FRS 1 expects the support being relied on to be specified.")
    if (not at_risk) and "material uncertainty" in full_text_low:
        contradictions.append(
            "The figures look solvent yet the notes flag a 'material uncertainty' over going "
            "concern — check this is consistent with the accounts.")

    base.update({
        "equity": equity, "current_assets": ca, "current_liabilities": cl,
        "net_current": net_current, "bs_insolvent": bs_insolvent,
        "liquidity_concern": liquidity_concern, "at_risk": at_risk,
        "financial_support": fin_support, "verdict": verdict,
        "verdict_level": verdict_level, "contradictions": contradictions,
    })
    if at_risk:
        base["elements"] = base["elements"] + [{
            "element": "Financial support specified (who provides it)",
            "present": fin_support,
        }]
    return base


# --------------------------------------------------------------------------
# AI review (Claude) — the judgement half: FRS compliance + grammar + summary.
# The deterministic checks above handle the arithmetic; this adds the reasoning.
# Requires the ANTHROPIC_API_KEY environment variable. Skips gracefully if unset.
# --------------------------------------------------------------------------
AI_MODEL = os.environ.get("FS_REVIEW_MODEL", "claude-haiku-4-5-20251001")
# Free option: Google Gemini. If GEMINI_API_KEY is set it is used (free tier);
# otherwise ANTHROPIC_API_KEY is used; otherwise the AI review is off.
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
# gemini-2.5-pro is NOT available on Google's free API tier (quota limit 0 -> 429),
# so default to gemini-2.5-flash, which IS free. Flash used to return empty output
# because "thinking" consumed the token budget; we now disable thinking (see
# _gemini_complete) so it returns full output. Override with GEMINI_MODEL if you
# have a paid key (e.g. gemini-2.5-pro). If ANTHROPIC_API_KEY is set and
# PREFER_CLAUDE=1, Claude is used instead.
GEMINI_MODEL = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")
PREFER_CLAUDE = os.environ.get("PREFER_CLAUDE", "0") == "1"
AI_ENABLED = bool(GEMINI_API_KEY or ANTHROPIC_API_KEY)


def _gemini_complete(prompt, key, max_tokens, json_out):
    import urllib.request, urllib.error
    url = ("https://generativelanguage.googleapis.com/v1beta/models/"
           + GEMINI_MODEL + ":generateContent")
    gen = {"maxOutputTokens": max_tokens, "temperature": 0}
    if json_out:
        gen["responseMimeType"] = "application/json"   # forces clean JSON, no fences
    # Disable "thinking" on 2.5 models so the whole token budget goes to the answer
    # (thinking previously ate the budget and returned empty output on flash).
    if "2.5" in GEMINI_MODEL:
        gen["thinkingConfig"] = {"thinkingBudget": 0}
    body = json.dumps({"contents": [{"parts": [{"text": prompt}]}],
                       "generationConfig": gen}).encode("utf-8")
    req = urllib.request.Request(url, data=body,
                                 headers={"Content-Type": "application/json",
                                          "x-goog-api-key": key})
    with urllib.request.urlopen(req, timeout=120) as r:
        data = json.loads(r.read().decode("utf-8"))
    cand = (data.get("candidates") or [{}])[0]
    parts = (cand.get("content") or {}).get("parts") or []
    text = "".join(p.get("text", "") for p in parts)
    if not text:
        # Empty usually means the token budget was consumed by "thinking".
        print(f"[ai_complete gemini] empty text; finishReason="
              f"{cand.get('finishReason')} usage={data.get('usageMetadata')}")
    return text or None


def _claude_complete(prompt, key, max_tokens):
    import anthropic
    client = anthropic.Anthropic(api_key=key)
    msg = client.messages.create(model=AI_MODEL, max_tokens=max_tokens,
                                 messages=[{"role": "user", "content": prompt}])
    return "".join(getattr(b, "text", "") for b in msg.content) or None


def ai_complete(prompt, max_tokens=8000, json_out=False):
    """Call the configured LLM and return its text, or None. Uses Claude when
    PREFER_CLAUDE=1 and a key is set; otherwise Gemini (2.5-pro by default);
    otherwise Anthropic as fallback. json_out asks Gemini for strict JSON."""
    order = []
    if PREFER_CLAUDE and ANTHROPIC_API_KEY:
        order = [("claude", ANTHROPIC_API_KEY)]
    elif GEMINI_API_KEY:
        order = [("gemini", GEMINI_API_KEY.strip())]
    elif ANTHROPIC_API_KEY:
        order = [("claude", ANTHROPIC_API_KEY)]
    for kind, key in order:
        try:
            if kind == "gemini":
                return _gemini_complete(prompt, key, max_tokens, json_out)
            return _claude_complete(prompt, key, max_tokens)
        except Exception as e:
            body = ""
            try:
                body = e.read().decode("utf-8", "replace")[:600]   # HTTPError body
            except Exception:
                pass
            print(f"[ai_complete {kind}] {type(e).__name__}: {e} {body}")
            return None
    return None


AI_PROMPT = """You are a senior technical reviewer of Singapore-incorporated \
companies' financial statements (SFRS / SFRS for Small Entities), the way a manager \
in an accounting firm reviews a junior's draft before it goes to the client. Be \
specific, cite the actual figures/notes from THIS set, and give practical, quotable \
recommendations. Do not pad with generic textbook statements; every observation must \
be anchored to something in this document.

A separate program has ALREADY checked the arithmetic; its findings are given to you \
below under "AUTOMATED CHECKS". Use them: comment on the accounting/disclosure \
IMPLICATIONS of any flagged item (e.g. a balance sheet that does not balance, cross-add \
errors in the statement of changes in equity, a going-concern risk), but do not merely \
re-list the numbers. You may still raise arithmetic points the program missed.

Review against the Singapore standards and report concrete issues:

FRS 1 / GOING CONCERN — treat this as a priority. Read the statement of financial \
position: is the company solvent (positive net assets AND net current assets, profitable) \
or not (net liabilities, net current liabilities, recurring losses, or negative operating \
cash flow)? If clearly solvent, a plain going-concern basis is fine — do NOT invent a \
material uncertainty. If there is ANY adverse indicator (net liabilities, net current \
liabilities, recurring losses, negative operating cash flow, or reliance on related-party \
funding), raise a HIGH-severity going-concern observation AND recommend a full disclosure \
covering EACH element below, tailored with this company's actual figures:
 (a) the adverse indicators, quoting the net-liability / net-current-liability / \
accumulated-loss figures;
 (b) whether a material uncertainty exists that may cast significant doubt on the ability \
to continue as a going concern, stated explicitly if so;
 (c) the specific financial support relied on — name the party (shareholder / director / \
holding company), quote the balance being supported (e.g. the amount due to that party), \
and state that a written undertaking / letter of support not to recall it for at least 12 \
months from the date of approval has been obtained;
 (d) management's assessment covering at least 12 months from the approval date, including \
cash-flow projections and funding availability;
 (e) the directors' conclusion that the going-concern basis remains appropriate;
 (f) the consequences if the company could not continue (assets written down to recoverable \
amount, non-current items reclassified as current, further liabilities provided for).
Flag whichever elements are missing, flag any contradiction between the narrative and the \
figures (e.g. a "can pay debts as they fall due" statement sitting oddly with net current \
liabilities), and ALWAYS include a ready-to-paste replacement going-concern note in \
suggested_wording whenever any indicator is present. Also check comparatives present and \
consistent; significant judgements & estimates note meaningful (not boilerplate); \
rounding/units stated; negatives in brackets consistently.

OTHER STANDARDS — be thorough and comprehensive; a competent reviewer of a small-company \
draft typically raises disclosure points across SEVERAL of these. For each applicable \
standard, check the specific disclosures below, and where any is missing or weak raise a \
separate observation with a concrete recommended disclosure:
 - FRS 1 presentation: statement of compliance with SFRS; basis of measurement; functional \
& presentation currency; material accounting-policy information; significant judgements and \
key estimation uncertainty (specific, not boilerplate); current/non-current split.
 - FRS 7 cash flows: correct operating / investing / financing classification; reconciliation \
of opening to closing cash that ties to the balance-sheet cash; non-cash transactions.
 - FRS 8: standards issued but not yet effective, with effective dates and expected impact; \
changes in estimates/policies and prior-period errors.
 - FRS 12 income taxes: current + deferred tax; a numerical effective-tax-rate \
reconciliation; the amount of unused tax losses / deductible temporary differences and \
whether a deferred tax asset is recognised or why not; the tax rate used.
 - FRS 16 PP&E: measurement basis; depreciation method and useful lives per class; \
reconciliation of carrying amount; the note's closing carrying amount MUST equal closing \
cost minus closing accumulated depreciation AND must equal the balance-sheet figure. If the \
note carrying amount differs from the balance sheet, the balance sheet (which balances) is \
usually the correct figure and the NOTE is the error — recommend correcting the note, NOT \
changing the balance sheet. Also check the depreciation charge in the note equals the \
depreciation in the profit & loss.
 - FRS 116 leases: right-of-use assets and lease liabilities for leased premises/equipment; \
maturity of lease liabilities; short-term / low-value exemptions — if the company clearly \
occupies premises but shows no lease or rent, say so.
 - FRS 115 revenue: over time vs point in time, stated and matching the ACTUAL business \
(flag any policy copied from a different industry); disaggregation of revenue; contract \
liabilities / deferred revenue where fees are prepaid.
 - FRS 109 financial instruments: the financial-instruments note must include ONLY financial \
instruments (exclude prepayments, statutory balances such as SDL/GST, suspense accounts, \
inventory and tax); categories; credit, liquidity (with a maturity analysis) and market \
risk; ECL / impairment approach for receivables.
 - FRS 24 related parties: name the relationships, disclose related-party transactions and \
outstanding balances with their terms, and disclose key management personnel compensation — \
flag if only a related-party definition/policy is given with no actual transactions / KMP note.
 - FRS 37 provisions, contingent liabilities and commitments; FRS 10 events after the \
reporting period; FRS 33 EPS if applicable; capital management; dividends.
 Also check directors'-statement consistency (directors' interests vs the share register), \
the company name and financial-year dates, and that every note cross-reference points to \
the correct note.

COMMON DRAFTING ERRORS — check each of these specifically, they recur in SME drafts:
 - INCOME TAX CONSISTENCY: the income tax expense on the face of the profit & loss MUST equal \
the TOTAL tax charge in the income-tax note, and the movement in the balance-sheet tax \
provision (opening + total charge − tax paid) must agree. IMPORTANT: an over- or \
under-provision in respect of prior years is a NORMAL, ACCEPTABLE component of the tax \
charge — total charge = current-year tax ± prior-year under/(over)-provision. So if the face \
figure differs from the note's current-year tax, first test whether the difference is a \
prior-year under/over-provision that makes the provision roll-forward tie; if it does, do \
NOT call the figures wrong — instead check the tax note discloses the under/(over)-provision \
as a separate line (e.g. "(Over)/under-provision in respect of prior year"), and recommend \
adding that line if missing. Only flag an ERROR when the figures cannot be reconciled even \
allowing for a prior-year provision adjustment. An over-provision credit should be shown in \
brackets.
 - LOANS TO / FROM DIRECTORS: any "amount due from director" (a loan TO a director) is a \
serious point — flag Companies Act 1967 section 162 (loans to directors are generally \
prohibited) and ask for confirmation of legality/exemption; for any material amount due \
to/from a director or shareholder, require FRS 24 disclosure of the terms (interest, \
security, repayment) and, for a receivable, a recoverability / ECL assessment. Note if the \
balance is a large proportion of total assets.
 - RETAINED EARNINGS TIE: retained earnings on the balance sheet MUST equal the closing \
retained earnings in the statement of changes in equity. If the balance sheet shows retained \
earnings before dividends with the dividend as a separate line while the SOCE shows it net, \
flag the presentation inconsistency.
 - DIVIDENDS: if a dividend is described as "payable" / declared but unpaid, it must appear \
as a liability on the balance sheet and must NOT be shown as a cash outflow in the financing \
section of the cash-flow statement; flag any inconsistency.
 - POLICY WITHOUT THE ASSET: if there is an accounting-policy or useful-life note for an \
asset class (e.g. property, plant and equipment) but that asset does not appear on the \
balance sheet, flag the redundant policy.
 - COMPARATIVES: every line and note must show the prior-year comparative; flag any line \
whose comparative is blank while the total column includes it.

SHARE CAPITAL — be precise: state the issued and paid-up share capital shown in the \
accounts as a dollar figure, check it is consistent across the statement of financial \
position, the statement of changes in equity and the share-capital note, and that it \
equals number of shares multiplied by issue price. If an ACRA figure appears in the \
automated checks above, say whether the two agree and by exactly how much, and give the \
most likely cause of any difference (e.g. the accounts showing the number of shares \
rather than the dollar amount, a share allotment after the year-end, or a transposition). \
Quote both figures.

Grammar & drafting: British spelling (SG uses British), singular/plural (Director vs \
Directors), defined-term capitalisation ("the Company"), leftover placeholders (square \
brackets, blanks, "XXX", TBD), wrong company name or financial-year dates, and \
inconsistent number formatting.

CORRECTED FIGURES — give the corrected presentation for numerical errors, but follow this \
ARITHMETIC DISCIPLINE strictly (you are a language model and must not rely on your own \
mental arithmetic):
 1. Use ONLY numbers that already appear in the financial statements or in the AUTOMATED \
CHECKS below. Do NOT invent or compute new totals in your head.
 2. To correct a figure, quote the two figures involved and state which one is right and \
what the other should change to — do not produce a long chain of new calculations.
 3. For a working-capital movement, take it directly as (this-year balance − last-year \
balance) using the two balance-sheet figures, and show that subtraction; a rise in a \
liability is a cash INFLOW (positive), a fall is an OUTFLOW (negative); the reverse for \
assets.
 4. HARD RULE for the statement of cash flows: the closing cash MUST equal the cash and \
cash equivalents on the balance sheet for the same year, and opening cash MUST equal last \
year's balance-sheet cash. State these two figures explicitly and make the statement tie to \
them. If your rebuilt operating/investing/financing lines do not add up to (closing − \
opening), DO NOT publish invented totals — instead list only the specific lines that are \
wrong and what each should be, and note that the full statement must be recast to close at \
the balance-sheet cash figure.
 5. Before finalising, re-read each corrected_figures block: if any total is one you \
calculated yourself rather than one taken from the statements/automated checks, remove the \
number and describe the correction in words instead. A correct qualitative correction beats \
a wrong number.

SUGGESTED WORDING — where a disclosure is missing, boilerplate, contradictory or wrong \
(e.g. an inadequate going-concern note, a revenue-recognition policy that describes the \
wrong business, a directors'-interest statement that contradicts the share register, a \
missing related-party note), draft the actual replacement paragraph the preparer can paste \
in, tailored to THIS company's figures and activities. Use square brackets only where a \
fact must be confirmed from the records.

Return STRICT JSON only (no prose, no markdown) in exactly this shape:
{"narrative":"3-6 sentence overall assessment a reviewer would write, naming the most \
important issues and the company's apparent financial position",
 "frs_observations":[{"area":"Going concern","frs":"FRS 1","severity":"high|medium|low",
   "issue":"short title","detail":"what is wrong/missing in THIS set, citing the figure or note",
   "recommendation":"what the preparer should do"}],
 "corrected_figures":[{"statement":"e.g. Statement of cash flows (FY2025)",
   "issue":"one line on what was wrong","corrected":"the corrected presentation, laid out \
line-by-line with \\n between lines, label then value","note":"why it now ties (optional)"}],
 "suggested_wording":[{"note":"e.g. Going concern (FRS 1)","draft":"the full replacement \
paragraph, ready to paste"}],
 "grammar_issues":[{"location":"where","current":"...","suggested":"..."}]}

Order frs_observations by severity (high first). Be comprehensive on FRS disclosures — work \
through every applicable standard above and raise a separate observation for each disclosure \
that is missing or weak (a real small-company draft usually has several across FRS 1, 7, 8, \
12, 16, 109, 115, 24 and 37). Whenever any going-concern indicator is present, include a \
detailed going-concern observation AND a ready-to-paste going-concern note in \
suggested_wording. Add corrected figures for every numerical error. Only omit a category if \
it is genuinely fine."""


def extract_full_text(doc):
    parts = []
    for i, table in enumerate(doc.tables):
        parts.append(f"\n--- TABLE {i + 1} ---")
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    return "\n".join(parts)


def _parse_json(raw):
    raw = (raw or "").strip()
    if raw.startswith("```"):
        raw = raw.split("```", 2)[1]
        if raw[:4].lower() == "json":
            raw = raw[4:]
    s, e = raw.find("{"), raw.rfind("}")
    if s >= 0 and e > s:
        try:
            return json.loads(raw[s:e + 1])
        except Exception:
            return {}
    return {}


def _findings_context(findings):
    """Summarise the deterministic checks as plain text to ground the AI review."""
    if not findings:
        return "(none provided)"
    L = []
    bc = findings.get("balance_checks") or []
    for b in bc:
        L.append(f"- Balance sheet: total assets {b['total_assets']:,.2f} vs equity+liabilities "
                 f"{b['equity_plus_liabilities']:,.2f} (diff {b['difference']:,.2f}) — "
                 f"{'balances' if b.get('balanced') else 'DOES NOT BALANCE'}.")
    for c in (findings.get("tally_checks") or [])[:12]:
        L.append(f"- Subtotal '{c['label']}' (table {c['table']}): lines sum to "
                 f"{c['sum_of_parts']:,.2f} but stated {c['stated_total']:,.2f} (diff {c['difference']:,.2f}).")
    for c in (findings.get("pl_checks") or [])[:8]:
        L.append(f"- P&L check '{c['check']}' (table {c['table']}): expected {c['expected']:,.2f}, "
                 f"stated {c['stated']:,.2f} (diff {c['difference']:,.2f}).")
    for c in (findings.get("row_checks") or [])[:12]:
        L.append(f"- Cross-add '{c['row']}' (table {c['table']}): across = {c['sum_across']:,.2f} "
                 f"vs stated total {c['stated_total']:,.2f} (diff {c['difference']:,.2f}).")
    for c in (findings.get("cross_checks") or []):
        L.append(f"- Cross-statement: {c['check']} — {c['left']:,.2f} vs {c['right']:,.2f} "
                 f"(diff {c['difference']:,.2f}).")
    for rp in (findings.get("related_party") or []):
        L.append(f"- Related party: {rp['error']} ({rp['recommendation']})")
    ca = findings.get("cash_anchor") or {}
    if ca.get("closing") is not None:
        if ca.get("opening") is not None:
            L.append(f"- GROUND-TRUTH cash (from the balance sheet — use these EXACT figures, "
                     f"do NOT recompute): closing cash = {ca['closing']:,.2f}, opening cash = "
                     f"{ca['opening']:,.2f}, so the net movement in cash MUST be "
                     f"{ca['net_change']:,.2f}. Any rebuilt statement of cash flows must close "
                     f"at {ca['closing']:,.2f}.")
        else:
            L.append(f"- GROUND-TRUTH cash (from the balance sheet): closing cash = "
                     f"{ca['closing']:,.2f}. Any rebuilt statement of cash flows must close at "
                     f"this figure.")
    gc = findings.get("going_concern") or {}
    if gc.get("verdict"):
        L.append(f"- Going concern: {gc['verdict']}")
    missing = [k["item"] for k in (findings.get("frs_checks") or []) if not k.get("present")]
    if missing:
        L.append("- FRS keyword scan did not find: " + "; ".join(missing) + ".")
    return "\n".join(L) if L else "No arithmetic/disclosure problems were flagged automatically."


def ai_review(extracted_text, findings=None):
    """Judgement-based review via the configured LLM, grounded in the automated checks."""
    if not AI_ENABLED:
        return {"enabled": False,
                "error": "AI review not enabled — set GEMINI_API_KEY or ANTHROPIC_API_KEY.",
                "frs_observations": [], "corrected_figures": [], "suggested_wording": [],
                "grammar_issues": [], "narrative": ""}
    prompt = (AI_PROMPT
              + "\n\nAUTOMATED CHECKS (already run — comment on their implications):\n"
              + _findings_context(findings)
              + "\n\nFINANCIAL STATEMENTS (extracted text follows):\n\n"
              + extracted_text[:70000])
    raw = ai_complete(prompt, max_tokens=12000, json_out=True)
    if not raw:
        return {"enabled": False,
                "error": "AI review could not run — check the API key (and any rate limits/credit).",
                "frs_observations": [], "corrected_figures": [], "suggested_wording": [],
                "grammar_issues": [], "narrative": ""}
    data = _parse_json(raw)
    obs = data.get("frs_observations", []) or []
    order = {"high": 0, "medium": 1, "low": 2}
    obs.sort(key=lambda o: order.get(str(o.get("severity", "")).lower(), 3))
    narrative = data.get("narrative", "")
    unverified = _unverified_ai_numbers(obs, extracted_text)
    if unverified:
        narrative = (narrative + "  ⚠ Please double-check these figures — the AI cited "
                     "them but they were not found verbatim in the document (they may be "
                     "sums the AI computed, which the free model can get wrong): "
                     + ", ".join(unverified) + ".").strip()
    return {
        "enabled": True, "error": None,
        "frs_observations": obs,
        "corrected_figures": data.get("corrected_figures", []) or [],
        "suggested_wording": data.get("suggested_wording", []) or [],
        "grammar_issues": data.get("grammar_issues", []) or [],
        "narrative": narrative,
    }


def _unverified_ai_numbers(observations, text):
    """Conservative hallucination guard: return the DISTINCT large figures (>= 1,000)
    that the AI quotes in its observation details but that do NOT appear (comma-
    stripped) anywhere in the document text. Small numbers, years and percentages are
    ignored to avoid false alarms."""
    text_digits = re.sub(r"[,\s]", "", text)
    years = {str(y) for y in range(1990, 2101)}
    seen, misses = set(), []
    for o in observations:
        blob = " ".join(str(o.get(k, "")) for k in ("detail", "issue", "recommendation"))
        for tok in re.findall(r"\d[\d,]{3,}", blob):        # 4+ digits (>= 1,000)
            norm = tok.replace(",", "")
            if norm in years or norm in seen:
                continue
            seen.add(norm)
            if norm not in text_digits:
                misses.append(tok)
    return misses[:8]


# --------------------------------------------------------------------------
# ACRA verification — checks the UEN against ACRA's free open data
# (data.gov.sg), and extracts the FS share capital for manual comparison
# against the (paid) ACRA Business Profile.
# --------------------------------------------------------------------------
import urllib.request
import urllib.parse

ACRA_RESOURCE = "d_3f960c10fed6145404ca7b821f263b87"
# Singapore UEN formats: 9-digit businesses, 10-char local companies (yyyy+5+letter),
# and the (T|S|R)yyXXnnnnX form for others.
UEN_RE = re.compile(
    r"\b(\d{9}[A-Z]|\d{8}[A-Z]|[TSR]\d{2}[A-Z]{2}\d{4}[A-Z])\b")


def _norm(s):
    return re.sub(r"[^a-z0-9]", "", (s or "").lower())


def extract_share_capital(doc):
    """The FS issued/paid-up share-capital dollar figure (SOFP or share-capital note).
    Matches the common label variants so the ACRA comparison is on the right line."""
    variants = ("issued and paid-up", "issued and paid up", "issued and fully paid",
                "called up share capital", "ordinary share capital",
                "at 1 january and 31 december", "at 31 december")
    candidates = []
    for table in doc.tables:
        labels, numgrid = _grid(table)
        skip = _note_columns(table)
        low_all = " ".join(labels).lower()
        note_is_share_capital = "share capital" in low_all or "ordinary share" in low_all
        for r, label in enumerate(labels):
            low = label.lower().strip()
            match = (low.startswith("share capital") or low == "share capital"
                     or any(v in low for v in variants[:5])
                     or (note_is_share_capital and any(v in low for v in variants[5:])))
            if not match:
                continue
            for c in range(1, len(numgrid[r])):
                if c in skip:
                    continue
                v = numgrid[r][c]
                if v is not None:
                    candidates.append(v)
    if not candidates:
        return None
    # Prefer the largest-magnitude candidate: the dollar amount of share capital,
    # not a stray note reference (e.g. "9") or a tiny mis-hit. A share-count column
    # equal to the dollar amount is harmless (same value).
    return max(candidates, key=lambda x: abs(x))


def acra_check(full_text):
    """Verify the UEN found in the document against ACRA's free open data."""
    out = {"enabled": True, "error": None, "uen": None, "found": False,
           "official_name": None, "status": None, "address": None,
           "name_matches": None}
    m = UEN_RE.search(full_text)
    if not m:
        out["enabled"] = False
        out["error"] = "No UEN / company registration number was found in the document."
        return out
    uen = m.group(1)
    out["uen"] = uen
    try:
        url = ("https://data.gov.sg/api/action/datastore_search?resource_id="
               + ACRA_RESOURCE + "&q=" + urllib.parse.quote(uen))
        req = urllib.request.Request(url, headers={"User-Agent": "fs-review"})
        with urllib.request.urlopen(req, timeout=15) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        records = data.get("result", {}).get("records", [])
        rec = next((x for x in records if x.get("uen", "").upper() == uen.upper()), None)
        if rec:
            out["found"] = True
            out["official_name"] = rec.get("entity_name")
            out["status"] = rec.get("uen_status_desc")
            out["address"] = (f"{rec.get('reg_street_name','')} "
                              f"{rec.get('reg_postal_code','')}").strip()
            out["name_matches"] = _norm(rec.get("entity_name")) in _norm(full_text)
        else:
            out["error"] = f"UEN {uen} was not found in the ACRA register."
    except Exception as e:
        out["enabled"] = False
        out["error"] = f"Could not reach the ACRA data service: {e}"
    return out


def crawl_bizfile(path):
    """Extract the official figures from an uploaded ACRA Business Profile PDF
    (UEN, name, status, issued/paid-up share capital, shareholders)."""
    out = {"error": None, "uen": None, "entity_name": None, "status": None,
           "issued_share_capital": None, "paid_up_capital": None,
           "shareholders": [], "directors": [], "ai": False}
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        text = "\n".join((p.extract_text() or "") for p in reader.pages)
    except Exception as e:
        out["error"] = f"Could not read the BizFile PDF: {e}"
        return out

    def find_amount(label):
        m = re.search(label + r"[^\d]{0,40}([\d,]+(?:\.\d+)?)", text, re.I)
        return _to_number(m.group(1)) if m else None

    out["issued_share_capital"] = find_amount(r"issued share capital")
    out["paid_up_capital"] = find_amount(r"paid[\s\-]?up capital")
    m = UEN_RE.search(text)
    if m:
        out["uen"] = m.group(1)

    if AI_ENABLED:
        prompt = (
            "This is the text of a Singapore ACRA Business Profile. Extract the "
            "official details as STRICT JSON only, in this shape: "
            '{"uen":"","entity_name":"","status":"","issued_share_capital":0,'
            '"paid_up_capital":0,"shareholders":[{"name":"","shares":0}],'
            '"directors":[{"name":""}]}. '
            "Use plain numbers (no commas/$) for capital and shares. Omit a field "
            "if not present. Text:\n\n" + text[:40000])
        raw = ai_complete(prompt, max_tokens=1500)
        data = _parse_json(raw) if raw else {}
        if data:
            out["ai"] = True
            for k in ("uen", "entity_name", "status"):
                if data.get(k):
                    out[k] = data[k]
            for k in ("issued_share_capital", "paid_up_capital"):
                v = data.get(k)
                if isinstance(v, (int, float)):
                    out[k] = float(v)
            if isinstance(data.get("shareholders"), list):
                out["shareholders"] = [
                    s for s in data["shareholders"] if isinstance(s, dict)][:20]
            if isinstance(data.get("directors"), list):
                out["directors"] = [
                    d for d in data["directors"] if isinstance(d, dict)][:20]
    return out


def check_cross_statements(doc):
    """Ties between the primary statements. Conservative: only flags when both
    sides are confidently located, to avoid false positives."""
    out = []

    def val(keys, exclude=()):
        for table in doc.tables:
            labels, numgrid = _grid(table)
            skip = _note_columns(table)
            ncols = max((len(r) for r in numgrid), default=0)
            for c in range(1, ncols):
                if c in skip:
                    continue
                v = _find_row(labels, numgrid, c, *keys, exclude=exclude)
                if v is not None:
                    return v
        return None

    # Statement of cash flows: opening + net movement should equal closing cash.
    end_cash = val(("cash and cash equivalents at end", "cash and cash equivalents at the end",
                    "cash and bank balances at end"))
    beg_cash = val(("cash and cash equivalents at beginning",
                    "cash and cash equivalents at the beginning",
                    "cash and bank balances at beginning"))
    net_cash = val(("net increase in cash", "net decrease in cash",
                    "net (decrease)/increase in cash", "net increase/(decrease) in cash",
                    "net change in cash"))
    if end_cash is not None and beg_cash is not None and net_cash is not None:
        if abs((beg_cash + net_cash) - end_cash) > 0.5:
            out.append({"check": "Cash flow — opening cash + net movement should equal closing cash",
                        "left": round(beg_cash + net_cash, 2), "right": round(end_cash, 2),
                        "difference": round((beg_cash + net_cash) - end_cash, 2)})
    # The cash-flow statement's closing cash MUST equal the balance-sheet cash.
    # Use the deterministic balance-sheet cash anchor (year columns only).
    _ca = cash_anchor(doc)
    bs_cash = _ca["closing"] if _ca else None
    if end_cash is not None and bs_cash is not None and abs(end_cash - bs_cash) > 0.5:
        out.append({"check": "Statement of cash flows closing cash does not tie to the "
                             "balance-sheet cash",
                    "left": round(end_cash, 2), "right": round(bs_cash, 2),
                    "difference": round(end_cash - bs_cash, 2)})
    return out


# Captions whose NOTE figure must tally back to the face of the P&L / balance
# sheet. (keys to match, labels to exclude). Income tax is deliberately absent —
# a face-vs-note tax difference can be a legitimate prior-year over/under-provision.
# (caption, matching keys, label excludes, max note-table rows or None).
# The row cap keeps P&L captions away from long breakdown statements.
NOTE_FACE_ITEMS = [
    ("Revenue", ("revenue",), ("cost",), 8),
    ("Cost of sales", ("cost of sales", "cost of goods"), (), 8),
    ("Administrative expenses",
     ("administrative expenses", "general and administrative"), (), 8),
    ("Cash and bank balances",
     ("cash and bank balance", "cash and cash equivalents", "cash at bank"),
     ("beginning", "end of", "net ", "flow"), None),
    ("Trade and other receivables",
     ("trade and other receivable", "trade receivable", "accounts receivable",
      "account receivable"), (), None),
    ("Trade and other payables",
     ("trade and other payable", "other payables", "accounts payable",
      "account payable"), (), None),
    ("Amount due from director", ("amount due from director",), (), None),
    ("Amount due to director", ("amount due to director",), (), None),
    ("Amount due from shareholder", ("amount due from shareholder",), (), None),
    ("Amount due to shareholder", ("amount due to shareholder",), (), None),
]


def _row_vals_unique(row):
    """Numeric values in a row, counting merged cells once, excluding column 0."""
    seen, vals = set(), []
    for i, cell in enumerate(row.cells):
        tc = id(cell._tc)
        if tc in seen:
            continue
        seen.add(tc)
        if i == 0:
            continue
        v = _to_number(cell.text)
        if v is not None:
            vals.append(v)
    return vals


def check_note_face_ties(doc):
    """Every note figure must tally back to the same caption on the face of the
    P&L / balance sheet (current year). Conservative, tested against real sets:
    - labels match by PREFIX and the note must match the SAME key as the face
      (so 'Other payables' is never compared against 'Trade and other payables');
    - note-reference columns are skipped and merged cells counted once;
    - cash-flow / SOCE tables are excluded (their movement lines share captions);
    - a note also ties if its unlabeled TOTAL row matches the face (a matched
      component line inside a breakdown is not a mismatch);
    - magnitudes are compared, so bracket conventions don't false-positive."""
    tables = list(doc.tables)

    def tbl_text(t):
        return " ".join(c.text for r in t.rows for c in r.cells).lower()

    BAD = ("increase", "decrease", "changes in", "movement")
    pl_idx = bs_idx = None
    skip_tables = set()
    for i, t in enumerate(tables):
        txt = tbl_text(t)
        if pl_idx is None and "before tax" in txt and ("revenue" in txt or "gross profit" in txt):
            pl_idx = i
        if bs_idx is None and ("total assets" in txt or "total equity and liabilities" in txt):
            bs_idx = i
        if "cash flows from" in txt:
            skip_tables.add(i)
    face_ids = {i for i in (pl_idx, bs_idx) if i is not None}
    if not face_ids:
        return []

    def vals_of(row, skipcols):
        seen, vals = set(), []
        for i, cell in enumerate(row.cells):
            tc = id(cell._tc)
            if tc in seen:
                continue
            seen.add(tc)
            if i == 0 or i in skipcols:
                continue
            v = _to_number(cell.text)
            if v is not None:
                vals.append(v)
        return vals

    def find_in(t, keys, excl):
        """Current-year value of the first row whose label matches a key (prefix)."""
        skipcols = _note_columns(t)
        for row in t.rows:
            label = (row.cells[0].text or "").lower().strip()
            if any(x in label for x in excl) or any(b in label for b in BAD):
                continue
            if any(label == k or label.startswith(k) for k in keys):
                vals = vals_of(row, skipcols)
                if vals:
                    return vals[-2] if len(vals) >= 2 else vals[-1]
        return None

    def table_total(t):
        """Current-year value of the table's unlabeled total row, if any."""
        skipcols = _note_columns(t)
        best = None
        for row in t.rows:
            label = (row.cells[0].text or "").strip()
            vals = vals_of(row, skipcols)
            if not label and vals:
                best = vals[-2] if len(vals) >= 2 else vals[-1]
        return best

    out = []
    for item, keys, excl, max_rows in NOTE_FACE_ITEMS:
        face = None
        for fi in (pl_idx, bs_idx):
            if fi is None:
                continue
            face = find_in(tables[fi], keys, excl)
            if face is not None:
                break
        if face is None:
            continue
        for j, t in enumerate(tables):
            if j in face_ids or j in skip_tables:
                continue
            if max_rows and len(t.rows) > max_rows:
                continue
            v = find_in(t, keys, excl)
            if v is None or abs(abs(v) - abs(face)) <= 0.5:
                continue
            tot = table_total(t)
            if tot is not None and abs(abs(tot) - abs(face)) <= 0.5:
                continue                      # the note's total ties — component match
            out.append({"check": f"Note-to-statement tie: '{item}' — a note shows "
                                 f"{abs(v):,.2f} but the statement shows {abs(face):,.2f}",
                        "left": round(abs(v), 2), "right": round(abs(face), 2),
                        "difference": round(abs(v) - abs(face), 2)})
            break
    return out


def cash_anchor(doc):
    """Deterministic ground-truth cash figures from the balance sheet, so the AI
    has the correct closing/opening cash and net movement to rebuild the cash-flow
    statement instead of computing them itself (where the free model slips)."""
    for table in doc.tables:
        labels, numgrid = _grid(table)
        low_all = " ".join(labels).lower()
        if "total current liabilit" not in low_all and "total asset" not in low_all:
            continue
        skip = _note_columns(table)
        ncols = max((len(r) for r in numgrid), default=0)
        moneycols = [c for c in range(1, ncols) if c not in skip]
        for r, label in enumerate(labels):
            low = label.lower().strip()
            if (low.startswith("cash and cash equivalents") or low == "cash"
                    or low.startswith("cash and bank")):
                vals = [numgrid[r][c] for c in moneycols
                        if c < len(numgrid[r]) and numgrid[r][c] is not None]
                # Year columns are the RIGHTMOST; a leftover note-reference column
                # (e.g. "8") sits on the left, so use the last two figures — current
                # year (closing) then prior year (opening).
                if len(vals) >= 2:
                    closing, opening = vals[-2], vals[-1]
                    return {"closing": round(closing, 2), "opening": round(opening, 2),
                            "net_change": round(closing - opening, 2)}
                if len(vals) == 1:
                    return {"closing": round(vals[-1], 2), "opening": None, "net_change": None}
    return None


def check_related_party_loans(doc):
    """Flag director/shareholder balances on the balance sheet. A loan TO a director
    is a Companies Act 1967 s.162 point; any material related-party balance needs FRS
    24 disclosure. Conservative: only scans the balance-sheet table (total assets /
    current liabilities present) so it doesn't double-fire on note tables."""
    out = []
    for table in doc.tables:
        labels, numgrid = _grid(table)
        low_all = " ".join(labels).lower()
        if "total current liabilit" not in low_all and "total asset" not in low_all:
            continue
        skip = _note_columns(table)
        ncols = max((len(r) for r in numgrid), default=0)
        moneycols = [c for c in range(1, ncols) if c not in skip]
        for r, label in enumerate(labels):
            low = label.lower().strip()
            vals = [numgrid[r][c] for c in moneycols
                    if c < len(numgrid[r]) and numgrid[r][c] is not None]
            amt = (vals[-2] if len(vals) >= 2 else vals[-1]) if vals else None
            if amt is None or abs(amt) < 100:      # ignore trivial/immaterial balances
                continue
            if any(k in low for k in ("due from director", "due from a director",
                                      "loan to director", "loan to a director",
                                      "loan to the director", "receivable from director")):
                out.append({"severity": "high", "kind": "loan to director",
                    "error": f"'{label.strip()}' of ${amt:,.2f} is an amount owed by a "
                             f"director to the company (a loan to a director).",
                    "recommendation": "Confirm this complies with Companies Act 1967 "
                             "section 162 (loans to directors are generally prohibited); "
                             "disclose the terms (interest, security, repayment) under FRS 24; "
                             "and assess its recoverability (expected credit loss)."})
            elif any(k in low for k in ("due from shareholder", "due from a shareholder",
                                        "due from related", "amount due from holding")):
                out.append({"severity": "medium", "kind": "due from related party",
                    "error": f"'{label.strip()}' of ${amt:,.2f} is a related-party receivable.",
                    "recommendation": "Disclose the terms (interest, security, repayment) under "
                             "FRS 24 and assess recoverability."})
            elif any(k in low for k in ("due to director", "due to a director",
                                        "due to shareholder", "due to a shareholder",
                                        "due to related", "amount due to holding")):
                out.append({"severity": "medium", "kind": "due to related party",
                    "error": f"'{label.strip()}' of ${amt:,.2f} is a related-party balance.",
                    "recommendation": "Disclose the nature and terms (interest, repayment, "
                             "security) under FRS 24."})
    seen, uniq = set(), []
    for o in out:
        k = (o["kind"], o["error"])
        if k not in seen:
            seen.add(k)
            uniq.append(o)
    return uniq


def build_disclosure_templates(findings, full_text_low=""):
    """Paste-ready Singapore FRS disclosure templates, triggered by what this review
    actually found and pre-filled with the company's own figures where known.
    Square brackets mark facts the preparer must confirm."""
    T = []

    gc = findings.get("going_concern") or {}
    if gc.get("at_risk"):
        eq = gc.get("equity")
        nc = gc.get("net_current")
        eq_s = f"net liabilities of ${abs(eq):,.0f}" if (eq is not None and eq < 0) \
            else "net assets of $[X]"
        nc_s = f"net current liabilities of ${abs(nc):,.0f}" if (nc is not None and nc < 0) \
            else "net current [assets/liabilities] of $[X]"
        T.append({"title": "Going concern (FRS 1)", "body": (
            "The Company incurred a [profit/loss] for the financial year of $[X] and, as at the "
            f"reporting date, had {nc_s} and {eq_s}. These conditions indicate that a material "
            "uncertainty exists that may cast significant doubt on the Company's ability to "
            "continue as a going concern.\n\n"
            "The financial statements have nevertheless been prepared on a going concern basis, "
            "the validity of which depends on the continued financial support of [the "
            "shareholder(s) / director(s) / holding company]. [Name] has provided a written "
            "undertaking not to demand repayment of the amount due to [him/her/it] of $[X] and "
            "to provide such financial support as is necessary for the Company to meet its "
            "obligations as and when they fall due, for at least twelve months from the date of "
            "approval of these financial statements. Management has prepared cash-flow "
            "projections covering that period.\n\n"
            "On this basis, the directors consider the use of the going concern basis to be "
            "appropriate. Should the Company be unable to continue as a going concern, "
            "adjustments would be required to write down assets to their recoverable amounts, "
            "reclassify non-current assets and liabilities as current, and provide for any "
            "further liabilities that might arise.")})

    rps = findings.get("related_party") or []
    if rps:
        lines = "\n".join("  - " + rp["error"] for rp in rps)
        T.append({"title": "Related party transactions (FRS 24)", "body": (
            "Related parties comprise the directors and shareholders of the Company and "
            "entities controlled by them.\n\n"
            "Significant balances with related parties at the reporting date:\n"
            f"{lines}\n"
            "These balances are unsecured, interest-free and repayable on demand[, except as "
            "disclosed in the going-concern note].\n\n"
            "Key management personnel compensation for the year:\n"
            "  Short-term employee benefits: $[X] (prior year: $[X]).\n\n"
            "Significant related-party transactions during the year: [describe — e.g. "
            "directors' remuneration, advances made/repaid, dividends declared].")})

    if any(rp.get("kind") == "loan to director" for rp in rps):
        T.append({"title": "Loan to a director — Companies Act 1967 s.162", "body": (
            "At the reporting date, an amount of $[X] was due from a director. The directors "
            "have assessed that this arrangement [falls within an exemption under section 162 "
            "of the Companies Act 1967 / was approved by the company in general meeting / was "
            "fully repaid on (date)]. The amount is [interest-free / interest-bearing at X% "
            "per annum], unsecured and repayable on demand.\n\n"
            "The directors have assessed the recoverability of the amount due from the "
            "director and are of the view that no allowance for expected credit losses is "
            "required [or: an allowance of $X has been recognised].")})

    if any(c.get("kind") == "tax sign / presentation"
           for c in (findings.get("pl_checks") or [])):
        T.append({"title": "Income tax expense (with prior-year over/under-provision)", "body": (
            "Income tax expense comprises:\n"
            "  Current year tax:                                    $[X]\n"
            "  (Over)/under-provision in respect of prior years:    $[X] / ($[X])\n"
            "  Income tax expense recognised in profit or loss:     $[X]\n\n"
            "The total tax expense reconciles to the profit before tax at the statutory rate "
            "of 17%, adjusted for the partial tax exemption, corporate income tax rebate, "
            "non-deductible expenses and [unutilised losses]. Note that the balance-sheet "
            "provision for taxation should roll forward: opening provision + total charge − "
            "tax paid = closing provision. An over-provision is a credit and is presented in "
            "brackets.")})

    if any(k.get("frs") == "FRS 116" and not k.get("present")
           for k in (findings.get("frs_checks") or [])):
        T.append({"title": "Leases (FRS 116) — choose the applicable option", "body": (
            "Option A — right-of-use asset recognised:\n"
            "The Company leases [premises] for its operations. A right-of-use asset and a "
            "lease liability are recognised at the lease commencement date. The right-of-use "
            "asset is depreciated on a straight-line basis over the lease term of [X] years; "
            "the lease liability is measured at the present value of the remaining lease "
            "payments, discounted at the Company's incremental borrowing rate of [X]%.\n\n"
            "Option B — exemption applied:\n"
            "The Company's lease of [premises] is a short-term lease with a lease term of 12 "
            "months or less [or: a lease of a low-value asset]. The Company has elected not to "
            "recognise a right-of-use asset and lease liability; lease payments are recognised "
            "as an expense on a straight-line basis over the lease term. Rent expense for the "
            "year was $[X] (prior year: $[X]).")})

    if gc.get("has_losses"):
        T.append({"title": "Unrecognised deferred tax assets (FRS 12)", "body": (
            "At the reporting date, the Company has unutilised tax losses of approximately "
            "$[X] (prior year: $[X]) available for offset against future taxable profits, "
            "subject to agreement by the Comptroller of Income Tax and compliance with the "
            "relevant provisions of the Income Tax Act 1947 (including the shareholding "
            "test). No deferred tax asset has been recognised in respect of these losses due "
            "to the uncertainty of the availability of future taxable profits.")})

    if "dividend" in full_text_low:
        T.append({"title": "Dividends", "body": (
            "During the financial year, the Company declared a [final/interim] one-tier "
            "tax-exempt dividend of $[X] ([X] cents per share) in respect of the financial "
            "year ended [date].\n\n"
            "[If unpaid at year-end:] The dividend remained unpaid at the reporting date and "
            "is presented as 'dividends payable' within current liabilities; it is NOT shown "
            "as a cash outflow in the statement of cash flows.\n"
            "[If paid during the year:] The dividend was paid during the year and is "
            "presented as a financing cash outflow in the statement of cash flows.")})

    return T


def build_corrections(findings):
    """Turn every finding into a concrete correction, split into the ERROR (what is
    wrong) and the RECOMMENDATION (what to do about it)."""
    C = []

    def add(sev, error, rec):
        C.append({"severity": sev, "error": error, "recommendation": rec})

    for b in findings.get("balance_checks", []):
        if not b.get("balanced"):
            add("high",
                f"ARITHMETIC — statement of financial position does not balance. · "
                f"Total assets stated: {b['total_assets']:,.2f} · "
                f"Total equity + liabilities: {b['equity_plus_liabilities']:,.2f} · "
                f"Difference to correct: {abs(b['difference']):,.2f}",
                "Recheck each subtotal on the balance sheet; total assets must equal "
                "total equity plus total liabilities.")
    for c in findings.get("pl_checks", []):
        if c.get("kind") == "tax sign / presentation":
            add("medium",
                f"Profit & loss (table {c['table']}): {c['check']}.",
                "A prior-year over/under-provision is acceptable — if so, show the credit "
                "in brackets and disclose an '(Over)/under-provision in respect of prior "
                "year' line in the tax note. Only if it is a genuine current-year expense "
                "does the profit/loss need restating.")
        else:
            add("high",
                f"ARITHMETIC — profit & loss does not cast (table {c['table']}, "
                f"{c['check']}). · Computed from the figures shown: {c['expected']:,.2f} · "
                f"Stated in the FS: {c['stated']:,.2f} · "
                f"Difference to correct: {abs(c['difference']):,.2f}",
                "Recompute this line — the error is usually in the tax line or a "
                "mis-cast subtotal.")
    for c in findings.get("tally_checks", []):
        add("high",
            f"ARITHMETIC — subtotal does not cast (table {c['table']}, '{c['label']}'). · "
            f"Sum of the line items above it: {c['sum_of_parts']:,.2f} · "
            f"Stated subtotal: {c['stated_total']:,.2f} · "
            f"Difference to correct: {abs(c['difference']):,.2f}",
            "Re-add the individual lines; either a line item or the stated subtotal "
            "is wrong.")
    for c in findings.get("row_checks", []):
        add("high",
            f"ARITHMETIC — row does not cast across (table {c['table']}, '{c['row']}'). · "
            f"Sum of the columns: {c['sum_across']:,.2f} · "
            f"Stated total column: {c['stated_total']:,.2f} · "
            f"Difference to correct: {abs(c['difference']):,.2f}",
            "Correct the row so the component columns add across to the total column.")
    for c in findings.get("cross_checks", []):
        add("high",
            f"ARITHMETIC — figures do not tie across statements ({c['check']}). · "
            f"One side: {c['left']:,.2f} · Other side: {c['right']:,.2f} · "
            f"Difference to correct: {abs(c['difference']):,.2f}",
            "These two figures describe the same thing and must be identical — "
            "correct whichever is wrong so they tie.")
    for rp in findings.get("related_party", []):
        add(rp["severity"], rp["error"], rp["recommendation"])
    gc = findings.get("going_concern", {})
    if gc.get("at_risk"):
        why = ("a net liabilities position" if gc.get("bs_insolvent")
               else "net current liabilities" if gc.get("liquidity_concern")
               else "losses / negative operating cash-flow indicators")
        missing = [e["element"] for e in gc.get("elements", []) if not e.get("present")]
        if missing:
            add("high",
                f"Going concern: the company shows {why}, but the note omits: {', '.join(missing)}.",
                "Strengthen the going-concern note to add and confirm the missing elements.")
        else:
            add("medium",
                f"Going concern: the company shows {why}.",
                "Confirm the note fully covers financial support, material uncertainty, the "
                "12-month assessment and the directors' conclusion.")
    for c in gc.get("contradictions", []):
        add("high", "Going concern contradiction: " + c,
            "Reconcile the going-concern narrative with the reported figures.")
    ac = findings.get("acra", {})
    fs_sc, reg = ac.get("fs_share_capital"), ac.get("registered_share_capital")
    if fs_sc is not None and reg is not None and not ac.get("share_capital_matches"):
        add("high",
            f"Share capital: the accounts show ${fs_sc:,.2f} but ACRA records ${reg:,.2f} "
            f"issued & paid-up (out by ${abs(reg - fs_sc):,.2f}).",
            "Reconcile the difference — check number of shares vs the dollar amount, or a "
            "post-year-end allotment.")
    if ac.get("found") and ac.get("name_matches") is False:
        add("medium",
            f"Company name differs from ACRA ('{ac.get('official_name')}').",
            "Correct the company name in the accounts to match ACRA.")
    for g in findings.get("language_issues", []):
        add("low",
            f"\"{g['found']}\" ({g.get('kind', 'wording')}).",
            f"Replace with \"{g['suggest']}\".")
    ai = findings.get("ai", {})
    # Going concern is already captured as a precise deterministic correction above
    # (it lists exactly which disclosure elements are missing), so skip the AI's
    # going-concern item here to avoid a near-duplicate row in the punch-list — its
    # full detail still appears in the FRS observations section.
    gc_covered = bool(gc.get("at_risk"))
    for o in (ai.get("frs_observations") or []):
        rec = o.get("recommendation")
        issue = o.get("issue") or o.get("detail")
        area = o.get("area") or o.get("frs") or "FRS"
        if gc_covered and "going concern" in (area + " " + (issue or "")).lower():
            continue
        if rec or issue:
            sev = str(o.get("severity", "medium")).lower()
            err = f"{area}: {issue}" if issue else f"{area}: {rec}"
            add(sev if sev in ("high", "medium", "low") else "medium",
                err, rec or "See the FRS observation above.")

    # Final safety net: drop any exact-duplicate (severity, error) rows.
    seen, uniq = set(), []
    for c in C:
        key = (c["severity"], c["error"].strip().lower())
        if key not in seen:
            seen.add(key)
            uniq.append(c)
    C = uniq
    order = {"high": 0, "medium": 1, "low": 2}
    C.sort(key=lambda x: order.get(x["severity"], 3))
    return C


def review_docx(path):
    """Return a dict of findings for a .docx file (rule-based, offline)."""
    findings = {
        "type": "docx",
        "sections_found": [], "sections_missing": [],
        "tables": 0, "paragraph_count": 0,
        "tally_checks": [], "balance_checks": [],
        "pl_checks": [], "row_checks": [], "cross_checks": [], "cash_anchor": None,
        "related_party": [], "disclosure_templates": [],
        "frs_checks": [], "language_issues": [],
        "ai": {"enabled": False, "error": None, "frs_observations": [],
               "corrected_figures": [], "suggested_wording": [],
               "grammar_issues": [], "narrative": ""},
        "acra": {"enabled": False, "error": None, "uen": None, "found": False,
                 "official_name": None, "status": None, "address": None,
                 "name_matches": None, "fs_share_capital": None,
                 "registered_share_capital": None, "share_capital_matches": None,
                 "bizfile": None},
        "going_concern": {"mentions_gc": False, "has_losses": False, "elements": []},
        "warnings": [], "error": None,
    }
    try:
        from docx import Document
    except Exception:
        findings["error"] = (
            "python-docx is not installed, so the .docx could not be parsed. "
            "Run: pip install python-docx"
        )
        return findings
    try:
        doc = Document(path)
    except Exception as e:
        findings["error"] = f"Could not open document: {e}"
        return findings

    full_text_low = "\n".join(p.text for p in doc.paragraphs).lower()
    # include table text for section / FRS keyword detection
    for table in doc.tables:
        for row in table.rows:
            full_text_low += "\n" + " ".join(c.text for c in row.cells).lower()

    findings["paragraph_count"] = len([p for p in doc.paragraphs if p.text.strip()])
    findings["tables"] = len(doc.tables)

    for section in KEY_SECTIONS:
        (findings["sections_found"] if section in full_text_low
         else findings["sections_missing"]).append(section.title())

    for t_idx, table in enumerate(doc.tables):
        findings["tally_checks"].extend(check_table_totals(t_idx, table))
        findings["pl_checks"].extend(check_pl(t_idx, table))
        findings["row_checks"].extend(check_row_totals(t_idx, table))

    findings["balance_checks"] = check_balance_equation(doc)
    findings["cross_checks"] = check_cross_statements(doc) + check_note_face_ties(doc)
    findings["cash_anchor"] = cash_anchor(doc)
    findings["related_party"] = check_related_party_loans(doc)
    findings["language_issues"] = check_language(doc)
    findings["frs_checks"] = check_frs(full_text_low, "inventor" in full_text_low)
    findings["going_concern"] = assess_going_concern(doc, full_text_low)
    # AI review runs last and is fed all the deterministic findings as grounding.
    findings["ai"] = ai_review(extract_full_text(doc), findings)

    acra = acra_check(extract_full_text(doc))
    acra["fs_share_capital"] = extract_share_capital(doc)
    acra["registered_share_capital"] = None
    acra["share_capital_matches"] = None
    findings["acra"] = acra

    findings["corrections"] = build_corrections(findings)
    findings["disclosure_templates"] = build_disclosure_templates(findings, full_text_low)

    findings["warnings"].append(
        "This review combines automated checks (arithmetic, balance-sheet and "
        "cross-statement ties, British-English spelling, an FRS disclosure checklist "
        "and a solvency-based going-concern test) with an AI FRS-judgement and grammar "
        "review. It is a first-pass aid — a qualified reviewer should still perform the "
        "final FRS/IFRS sign-off."
    )
    return findings


def basic_review(path, ext):
    if ext == ".docx":
        return review_docx(path)
    return {
        "type": ext.lstrip("."),
        "error": None,
        "warnings": [
            f"Automated parsing for {ext} files is not enabled in this build. "
            "Please upload the financial statements as a .docx for the full "
            "automated review. The file has been stored."
        ],
        "sections_found": [], "sections_missing": [],
        "tables": 0, "paragraph_count": 0,
        "tally_checks": [], "balance_checks": [],
        "pl_checks": [], "row_checks": [],
        "frs_checks": [], "language_issues": [],
        "ai": {"enabled": False, "error": None, "frs_observations": [],
               "corrected_figures": [], "suggested_wording": [],
               "grammar_issues": [], "narrative": ""},
        "acra": {"enabled": False, "error": None, "uen": None, "found": False,
                 "official_name": None, "status": None, "address": None,
                 "name_matches": None, "fs_share_capital": None,
                 "registered_share_capital": None, "share_capital_matches": None,
                 "bizfile": None},
        "going_concern": {"mentions_gc": False, "has_losses": False, "elements": []},
    }


# --------------------------------------------------------------------------
# App routes
# --------------------------------------------------------------------------
@app.route("/")
@login_required
def dashboard():
    records = list_records()
    return render_template("dashboard.html", records=records, name=session.get("name"))


@app.route("/upload", methods=["POST"])
@login_required
def upload():
    file = request.files.get("file")
    if not file or file.filename == "":
        flash("Please choose a file to upload.", "error")
        return redirect(url_for("dashboard"))

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        flash(f"File type {ext or '(none)'} is not allowed. "
              f"Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}.", "error")
        return redirect(url_for("dashboard"))

    rec_id = uuid.uuid4().hex[:12]
    safe_name = secure_filename(file.filename)
    stored_name = f"{rec_id}_{safe_name}"
    stored_path = os.path.join(UPLOAD_DIR, stored_name)
    file.save(stored_path)

    findings = basic_review(stored_path, ext)

    # Optional: crawl an uploaded ACRA Business Profile (BizFile) PDF and
    # cross-check its share capital / UEN against the financial statements.
    bizfile = request.files.get("acra_bizfile")
    if bizfile and bizfile.filename and "acra" in findings:
        bz_name = secure_filename(bizfile.filename)
        bz_path = os.path.join(UPLOAD_DIR, f"{rec_id}_bizfile_{bz_name}")
        bizfile.save(bz_path)
        bz = crawl_bizfile(bz_path)
        findings["acra"]["bizfile"] = bz
        reg = bz.get("paid_up_capital")
        if reg is None:
            reg = bz.get("issued_share_capital")
        findings["acra"]["registered_share_capital"] = reg
        fs_sc = findings["acra"].get("fs_share_capital")
        if reg is not None and fs_sc is not None:
            findings["acra"]["share_capital_matches"] = abs(reg - fs_sc) <= 0.5
        # Rebuild the correction punch-list now the ACRA share-capital comparison is in.
        findings["corrections"] = build_corrections(findings)

    with open(stored_path, "rb") as _fb:
        file_bytes = _fb.read()

    record = {
        "id": rec_id,
        "original_name": file.filename,
        "stored_name": stored_name,
        "ext": ext,
        "size_bytes": os.path.getsize(stored_path),
        "uploaded_by": session.get("name"),
        "uploaded_at": dt.datetime.now().isoformat(timespec="seconds"),
        "findings": findings,
    }
    save_record(record, file_bytes)
    log_event("upload & review", f"{file.filename} ({rec_id})")

    flash("File uploaded and reviewed.", "success")
    return redirect(url_for("report", rec_id=rec_id))


@app.route("/report/<rec_id>")
@login_required
def report(rec_id):
    record = get_record(rec_id)
    if not record:
        abort(404)
    return render_template("report.html", r=record)


@app.route("/download/<rec_id>")
@login_required
def download(rec_id):
    from flask import send_file
    got = get_record_file(rec_id)
    if not got:
        abort(404)
    name, data = got
    log_event("download original", f"{name} ({rec_id})")
    return send_file(io.BytesIO(data), as_attachment=True, download_name=name)


@app.route("/delete/<rec_id>", methods=["POST"])
@login_required
def delete(rec_id):
    delete_record(rec_id)
    log_event("delete record", rec_id)
    flash("Record deleted.", "success")
    return redirect(url_for("dashboard"))


def build_word_report(record):
    """Build a .docx review report from a record's findings; returns a BytesIO."""
    import io
    from docx import Document as _Doc
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    DARK, RED, GREEN, AMBER = "1F3864", "FCE4D6", "E2EFDA", "FFF2CC"
    f = record["findings"]

    def setfont(run, size=10, bold=False, color=None):
        run.font.name = "Arial"; run.font.size = Pt(size); run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

    def shade(cell, hexc):
        tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
        sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexc); tcPr.append(sh)

    def H(text, size=13):
        p = doc.add_paragraph(); setfont(p.add_run(text), size, True, DARK)

    def body(text, size=10, bold=False):
        p = doc.add_paragraph(); setfont(p.add_run(text), size, bold)

    def table(headers, rows, shades=None):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]; c.text = ""
            setfont(c.paragraphs[0].add_run(h), 9, True, "FFFFFF"); shade(c, DARK)
        for ri, row in enumerate(rows):
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = ""; setfont(cells[i].paragraphs[0].add_run(str(v)), 9)
                if shades and shades[ri]:
                    shade(cells[i], shades[ri])

    doc = _Doc()
    doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(10)
    # Firm branding: logo if logo.png is present, wordmark otherwise.
    if os.path.exists(LOGO_PATH):
        from docx.shared import Inches as _Inches
        lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.add_run().add_picture(LOGO_PATH, width=_Inches(1.8))
    else:
        bp = doc.add_paragraph(); bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        setfont(bp.add_run("ASSEMBLY WORKS"), 13, True, "2E5496")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    setfont(p.add_run("FINANCIAL STATEMENTS REVIEW REPORT"), 17, True, DARK)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    setfont(p.add_run(record["original_name"]), 11, True, "2E5496")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    setfont(p.add_run("Reviewed " + record["uploaded_at"].replace("T", " ")), 9, False, "606060")

    if f.get("error"):
        body(f["error"]);
    else:
        ac = f.get("acra", {})
        if ac.get("uen"):
            H("ACRA verification")
            rows = [["UEN in document", ac.get("uen", "")]]
            if ac.get("found"):
                nm = (ac.get("official_name") or "") + (
                    "  (matches)" if ac.get("name_matches") else "  (differs)")
                rows.append(["ACRA registered name", nm])
                rows.append(["Registration status", ac.get("status", "")])
                rows.append(["Registered address", ac.get("address", "")])
            elif ac.get("error"):
                rows.append(["Result", ac.get("error")])
            sc = ac.get("fs_share_capital")
            rows.append(["Share capital (per FS)",
                         f"{sc:,.2f}" if sc is not None else "not detected"])
            rsc = ac.get("registered_share_capital")
            if rsc is not None:
                rows.append(["Share capital (per ACRA)", f"{rsc:,.2f}" + (
                    "  (matches)" if ac.get("share_capital_matches") else "  (mismatch)")])
            table(["Field", "Value"], rows)

        if f.get("corrections"):
            H("Corrections to make")
            body("Each row shows the error found and the recommended correction.")
            sevlabel = {"high": "High", "medium": "Medium", "low": "Minor"}
            table(["#", "Severity", "Error found", "Recommended correction"],
                  [[str(i), sevlabel.get(c.get("severity", ""), ""),
                    c.get("error", c.get("text", "")), c.get("recommendation", "")]
                   for i, c in enumerate(f["corrections"], 1)],
                  [RED if c.get("severity") == "high"
                   else AMBER if c.get("severity") == "medium" else None
                   for c in f["corrections"]])

        H("Numerical & arithmetic findings")
        if f["tally_checks"] or f["pl_checks"] or f["row_checks"] or f.get("cross_checks") or \
           any(not b["balanced"] for b in f["balance_checks"]):
            rows, sh = [], []
            for c in f["tally_checks"]:
                rows.append([f"Table {c['table']}", c["label"],
                             f"{c['sum_of_parts']:,.2f}", f"{c['stated_total']:,.2f}",
                             c.get("kind", "sum")]); sh.append(RED)
            for c in f["pl_checks"]:
                rows.append([f"Table {c['table']}", c["check"],
                             f"{c['expected']:,.2f}", f"{c['stated']:,.2f}", "P&L"]); sh.append(RED)
            for c in f["row_checks"]:
                rows.append([f"Table {c['table']}", c["row"],
                             f"{c['sum_across']:,.2f}", f"{c['stated_total']:,.2f}", "cross-add"]); sh.append(RED)
            for b in f["balance_checks"]:
                if not b["balanced"]:
                    rows.append(["Balance sheet", "Assets vs Equity+Liabilities",
                                 f"{b['total_assets']:,.2f}", f"{b['equity_plus_liabilities']:,.2f}",
                                 "balance"]); sh.append(RED)
            for c in f.get("cross_checks", []):
                rows.append(["Cross-statement", c["check"],
                             f"{c['left']:,.2f}", f"{c['right']:,.2f}", "tie"]); sh.append(RED)
            table(["Source", "Item", "Calculated", "Reported", "Type"], rows, sh)
        else:
            body("All arithmetic checks passed (totals, balance equation, P&L flow, "
                 "cross-adds and cross-statement ties).")

        ai = f.get("ai", {})
        if ai.get("enabled"):
            if ai.get("narrative"):
                H("Reviewer's assessment"); body(ai["narrative"])
            if ai.get("frs_observations"):
                H("FRS compliance observations")
                table(["Severity", "Area / FRS", "Issue & detail", "Recommendation"],
                      [[(o.get("severity", "") or "").title(),
                        ((o.get("area", "") or "") + (" (" + o.get("frs", "") + ")" if o.get("frs") else "")).strip(),
                        (o.get("issue", "") or "") + ((" — " + o.get("detail", "")) if o.get("detail") else ""),
                        o.get("recommendation", "")] for o in ai["frs_observations"]])
            if ai.get("corrected_figures"):
                H("Corrected figures — ready to paste")
                for c in ai["corrected_figures"]:
                    body(c.get("statement", ""), 11, True)
                    if c.get("issue"):
                        body(c["issue"], 9)
                    for line in (c.get("corrected", "") or "").split("\n"):
                        body(line, 10)
                    if c.get("note"):
                        body(c["note"], 9)
            if ai.get("suggested_wording"):
                H("Suggested replacement wording")
                for s in ai["suggested_wording"]:
                    body(s.get("note", ""), 11, True)
                    body(s.get("draft", ""), 10)
            if ai.get("grammar_issues"):
                H("Grammar & wording")
                table(["Location", "Current", "Suggested"],
                      [[g.get("location", ""), g.get("current", ""), g.get("suggested", "")]
                       for g in ai["grammar_issues"]])
        else:
            H("AI review"); body(ai.get("error", "AI review not enabled."))

        if f.get("disclosure_templates"):
            H("Disclosure templates (paste-ready)")
            body("Triggered by this review's findings — adapt the [bracketed] items to the "
                 "company's facts before pasting into the notes.", 9)
            for t in f["disclosure_templates"]:
                body(t.get("title", ""), 11, True)
                for line in (t.get("body", "") or "").split("\n"):
                    body(line, 10)

        H("FRS disclosure checklist")
        table(["FRS", "Disclosure", "Detected"],
              [[k["frs"], k["item"], "Yes" if k["present"] else "Check"] for k in f["frs_checks"]],
              [GREEN if k["present"] else RED for k in f["frs_checks"]])

        gc = f.get("going_concern", {})
        if gc.get("elements"):
            H("Going concern (FRS 1 template)")
            table(["Element", "In the note?"],
                  [[e["element"], "Yes" if e["present"] else "Check"] for e in gc["elements"]],
                  [GREEN if e["present"] else RED for e in gc["elements"]])

    p = doc.add_paragraph()
    setfont(p.add_run("This automated review is a first-pass aid, not a substitute for a full "
                      "FRS/IFRS compliance review by a qualified reviewer."), 8, False, "808080")
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf


def _inject_comments_part(docx_bytes, comments):
    """Post-process a .docx: add word/comments.xml (+ rels + content-type) so the
    w:commentReference anchors placed in the document resolve to real Word margin
    comments. `comments` is a list of (id, text)."""
    import zipfile
    from xml.sax.saxutils import escape
    stamp = dt.datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
    items = []
    for cid, text in comments:
        paras = "".join(
            f"<w:p><w:r><w:t xml:space=\"preserve\">{escape(line)}</w:t></w:r></w:p>"
            for line in (text or "").split("\n"))
        items.append(f'<w:comment w:id="{cid}" w:author="FS Review" '
                     f'w:date="{stamp}" w:initials="FSR">{paras}</w:comment>')
    comments_xml = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                    '<w:comments xmlns:w="http://schemas.openxmlformats.org/'
                    'wordprocessingml/2006/main">' + "".join(items) + "</w:comments>")
    override = (b'<Override PartName="/word/comments.xml" ContentType="application/'
                b'vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>')
    rel = (b'<Relationship Id="rIdFsRevCmts" Type="http://schemas.openxmlformats.org/'
           b'officeDocument/2006/relationships/comments" Target="comments.xml"/>')
    zin = zipfile.ZipFile(io.BytesIO(docx_bytes))
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in zin.namelist():
            data = zin.read(name)
            if name == "[Content_Types].xml" and b"comments+xml" not in data:
                data = data.replace(b"</Types>", override + b"</Types>")
            elif name == "word/_rels/document.xml.rels" and b"relationships/comments" not in data:
                data = data.replace(b"</Relationships>", rel + b"</Relationships>")
            zout.writestr(name, data)
        zout.writestr("word/comments.xml", comments_xml)
    out.seek(0)
    return out


def build_marked_fs(record, file_bytes):
    """Revised FS module: regenerate the ORIGINAL financial statements .docx as a
    reviewed draft —
      - every review point is a real Word margin COMMENT (author "FS Review"),
        anchored at the relevant statement/table, for staff to work through and
        resolve in Word's Review pane;
      - unambiguous text corrections are applied as tracked changes staff can
        Accept/Reject;
      - proposed additional disclosures are appended in blue for insertion."""
    from docx import Document as _Doc
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = _Doc(io.BytesIO(file_bytes))
    f = record["findings"]
    AUTHOR = "FS Review"
    STAMP = dt.datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
    _rid = [9001]

    def _tracked_replace(par, old, new):
        """Replace old->new in a paragraph as a Word tracked change (accept/reject).

        Handles phrases that SPAN MULTIPLE RUNS — Word routinely splits sentences
        into many runs (spell-check, formatting), which previously made phrase
        corrections fall back to comments instead of being marked up."""
        import copy as _copy
        if not old:
            return False
        runs = par.runs
        full = "".join(r.text or "" for r in runs)
        idx = full.find(old)
        if idx < 0:
            return False
        end = idx + len(old)
        pos, affected = 0, []
        for r in runs:
            t = r.text or ""
            rs, re_ = pos, pos + len(t)
            if rs < end and re_ > idx:
                affected.append((r, max(idx, rs) - rs, min(end, re_) - rs))
            pos = re_
        if not affected:
            return False
        single = len(affected) == 1
        r0, s0, e0 = affected[0]
        t0 = r0.text or ""
        suffix_single = t0[e0:] if single else None
        r0.text = t0[:s0]
        if not single:
            for (r, s, e) in affected[1:-1]:
                r.text = ""
            rl, sl, el = affected[-1]
            rl.text = (rl.text or "")[el:]
        rpr = r0._element.find(qn("w:rPr"))
        anchor = r0._element
        d = OxmlElement("w:del")
        d.set(qn("w:id"), str(_rid[0])); _rid[0] += 1
        d.set(qn("w:author"), AUTHOR); d.set(qn("w:date"), STAMP)
        dr = OxmlElement("w:r")
        if rpr is not None:
            dr.append(_copy.deepcopy(rpr))
        dtx = OxmlElement("w:delText")
        dtx.set(qn("xml:space"), "preserve"); dtx.text = old
        dr.append(dtx); d.append(dr)
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), str(_rid[0])); _rid[0] += 1
        ins.set(qn("w:author"), AUTHOR); ins.set(qn("w:date"), STAMP)
        ir = OxmlElement("w:r")
        if rpr is not None:
            ir.append(_copy.deepcopy(rpr))
        it = OxmlElement("w:t")
        it.set(qn("xml:space"), "preserve"); it.text = new
        ir.append(it); ins.append(ir)
        anchor.addnext(ins); anchor.addnext(d)
        if single and suffix_single:
            tr = OxmlElement("w:r")
            if rpr is not None:
                tr.append(_copy.deepcopy(rpr))
            tt = OxmlElement("w:t")
            tt.set(qn("xml:space"), "preserve"); tt.text = suffix_single
            tr.append(tt); ins.addnext(tr)
        return True

    # ---- Word margin-comment machinery -------------------------------------
    _comments = []          # (id, text) for the comments part
    _next_cid = [0]

    def _anchor_comment(par, text):
        """Attach a margin comment to a paragraph (wraps the whole paragraph)."""
        cid = _next_cid[0]; _next_cid[0] += 1
        p = par._p
        start = OxmlElement("w:commentRangeStart"); start.set(qn("w:id"), str(cid))
        end = OxmlElement("w:commentRangeEnd"); end.set(qn("w:id"), str(cid))
        ref_r = OxmlElement("w:r")
        ref = OxmlElement("w:commentReference"); ref.set(qn("w:id"), str(cid))
        ref_r.append(ref)
        pPr = p.find(qn("w:pPr"))
        if pPr is not None:
            pPr.addnext(start)
        else:
            p.insert(0, start)
        p.append(end)
        p.append(ref_r)
        _comments.append((cid, text))

    def _find_paragraph(*needles, min_len=0):
        """First paragraph (body or table cell) whose text contains any needle."""
        needles = [n.lower() for n in needles if n]
        for par in doc.paragraphs:
            low = (par.text or "").lower()
            if len(low) >= min_len and any(n in low for n in needles):
                return par
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        low = (par.text or "").lower()
                        if len(low) >= min_len and any(n in low for n in needles):
                            return par
        return None

    def _table_anchor(t_idx):
        """First non-empty paragraph inside table t_idx (1-based)."""
        if t_idx and 0 < t_idx <= len(doc.tables):
            for row in doc.tables[t_idx - 1].rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        if (par.text or "").strip():
                            return par
        return None

    # --- 1) Tracked-change text corrections (typos + AI grammar suggestions) ---
    fixes = [(g.get("found"), g.get("suggest")) for g in (f.get("language_issues") or [])]
    fixes += [(g.get("current"), g.get("suggested"))
              for g in ((f.get("ai") or {}).get("grammar_issues") or [])]
    applied = 0
    for old, new in fixes:
        if not old or not new or old == new or len(old) > 300:
            continue
        for par in doc.paragraphs:
            if _tracked_replace(par, old, new):
                applied += 1
                break
        else:
            for tbl in doc.tables:
                done = False
                for row in tbl.rows:
                    for cell in row.cells:
                        for par in cell.paragraphs:
                            if _tracked_replace(par, old, new):
                                applied += 1; done = True; break
                        if done: break
                    if done: break
                if done: break

    # --- 2) Reviewer comments anchored under the flagged tables ---
    # ---- Every review point becomes a margin comment at the right spot -----
    import re as _re
    sevlabel = {"high": "HIGH", "medium": "MEDIUM", "low": "Minor"}
    first_par = next((p for p in doc.paragraphs if (p.text or "").strip()), None)

    # Map disclosure templates so the matching one is embedded INSIDE the comment —
    # the accountant copies the wording straight from the comment bubble.
    tmap = {}
    for tdef in (f.get("disclosure_templates") or []):
        tl = (tdef.get("title") or "").lower()
        if "loan to a director" in tl:
            tmap["s162"] = tdef
        elif "related party" in tl:
            tmap["rp"] = tdef
        elif "going concern" in tl:
            tmap["gc"] = tdef
        elif "income tax" in tl:
            tmap["tax"] = tdef

    def _template_for(low):
        if "loan to a director" in low or "owed by a director" in low or "s.162" in low \
                or "section 162" in low:
            return tmap.get("s162") or tmap.get("rp")
        if "related-party" in low or "related party" in low or "due to shareholder" in low \
                or "due from shareholder" in low or "due to director" in low \
                or "due from director" in low:
            return tmap.get("rp")
        if "going concern" in low:
            return tmap.get("gc")
        if "added back" in low or "over-provision" in low or "under-provision" in low:
            return tmap.get("tax")
        return None

    for i, c in enumerate((f.get("corrections") or []), 1):
        err = c.get("error", "") or ""
        rec = c.get("recommendation", "") or ""
        # " · " separates labelled figures (Computed / Stated / Difference) — give
        # each its own line in the comment bubble so the arithmetic reads clearly.
        err_lines = err.replace(" · ", "\n")
        text = (f"Assembly Works — Review point {i} "
                f"[{sevlabel.get(c.get('severity'), '')}]:\n{err_lines}\n"
                f"Action: {rec}")
        low = err.lower()
        tdef = _template_for(low)
        if tdef:
            text += ("\n\nSuggested disclosure — " + tdef.get("title", "") +
                     " (adapt the [bracketed] items):\n" + (tdef.get("body") or ""))
        anchor = None
        mt = _re.search(r"table (\d+)", low)
        if mt:
            anchor = _table_anchor(int(mt.group(1)))
        if anchor is None and "going concern" in low:
            anchor = _find_paragraph("going concern")
        if anchor is None and ("due from director" in low or "due to director" in low
                               or "due from shareholder" in low or "due to shareholder" in low
                               or "related-party" in low or "related party" in low):
            anchor = _find_paragraph("due from director", "due to director",
                                     "due from shareholder", "due to shareholder",
                                     "due to related", "amount due")
        if anchor is None and ("does not balance" in low or "balance sheet" in low
                               or "financial position" in low):
            anchor = _find_paragraph("total equity and liabilities", "total assets")
        if anchor is None and "cash" in low:
            anchor = _find_paragraph("cash flows from operating", "statement of cash flows",
                                     "cash and cash equivalents")
        if anchor is None and "share capital" in low:
            anchor = _find_paragraph("share capital")
        if anchor is None and "tax" in low:
            anchor = _find_paragraph("income tax", "taxation")
        if anchor is None:
            anchor = first_par
        if anchor is not None:
            _anchor_comment(anchor, text)

    # Pointer comment for the appended disclosure templates.
    if f.get("disclosure_templates") and first_par is not None:
        names = "; ".join(t.get("title", "") for t in f["disclosure_templates"])
        _anchor_comment(first_par,
                        "FS Review: proposed additional disclosures are appended at the end "
                        "of this document (in blue) — insert into the notes and adapt the "
                        "[bracketed] items. Templates: " + names + ".")

    # --- 4) Proposed additional disclosures, appended as insertions ---
    if f.get("disclosure_templates"):
        hp = doc.add_paragraph()
        r = hp.add_run("PROPOSED ADDITIONAL DISCLOSURES (INSERT — adapt [bracketed] items)")
        r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor.from_string("1F3864")
        for tdef in f["disclosure_templates"]:
            tp = doc.add_paragraph()
            r = tp.add_run(tdef.get("title", "")); r.bold = True
            r.font.color.rgb = RGBColor.from_string("2E5496"); r.font.size = Pt(11)
            for line in (tdef.get("body", "") or "").split("\n"):
                bp = doc.add_paragraph()
                br = bp.add_run(line)
                br.font.color.rgb = RGBColor.from_string("2E5496")
                br.font.size = Pt(10)
                br.underline = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    if _comments:
        return _inject_comments_part(buf.read(), _comments)
    return buf


@app.route("/report/<rec_id>/revised.docx")
@login_required
def download_revised(rec_id):
    """Download the original FS with the review's recommendations marked up."""
    from flask import send_file
    record = get_record(rec_id)
    got = get_record_file(rec_id)
    if not record or not got:
        abort(404)
    name, data = got
    if not name.lower().endswith(".docx"):
        abort(400)
    buf = build_marked_fs(record, data)
    log_event("download revised FS", f"{record['original_name']} ({rec_id})")
    out = os.path.splitext(record["original_name"])[0] + "_revised (marked up).docx"
    return send_file(buf, as_attachment=True, download_name=out,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/report/<rec_id>/report.docx")
@login_required
def download_report(rec_id):
    from flask import send_file
    record = get_record(rec_id)
    if not record:
        abort(404)
    buf = build_word_report(record)
    log_event("download review report", f"{record['original_name']} ({rec_id})")
    name = os.path.splitext(record["original_name"])[0] + "_reviewed.docx"
    return send_file(buf, as_attachment=True, download_name=name,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.template_filter("filesize")
def filesize(n):
    for unit in ["B", "KB", "MB", "GB"]:
        if n < 1024:
            return f"{n:.0f} {unit}" if unit == "B" else f"{n:.1f} {unit}"
        n /= 1024
    return f"{n:.1f} TB"


# Initialise storage (create tables + seed admin) at startup. Wrapped so a
# transient DB hiccup doesn't stop the app from booting.
try:
    init_db()
except Exception as _e:
    print(f"[init_db] warning: {_e}")


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
